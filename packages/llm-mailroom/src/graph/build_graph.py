import functools
import structlog
import threading
import time
import uuid
from pathlib import Path
from typing import Any, Mapping

from langgraph.graph import StateGraph, START, END
from langgraph.checkpoint.memory import MemorySaver
from langgraph.types import Command, interrupt

from graph.state import DocumentState
from graph.routing import (
    after_classify,
    after_retry_classify,
    after_review_classify,
    after_extraction,
    after_extraction_gated,
    after_retry_extraction,
    after_retry_extraction_gated,
    after_judge,
    after_arbiter,
    after_boss,
    after_human_review,
    after_report,
    judge_gate,
)
from observability.tracing import pipeline_trace, traced_node, observation
from schemas.manifest import DocumentManifest, PipelineStage
from pipeline.config import get_confidence_thresholds, UNKNOWN_DOC_TYPE
from pipeline.bins import (
    inbox_dir,
    processing_dir,
    classified_dir,
    review_dir,
    failed_dir,
    archive_dir,
    manifests_dir,
    ensure_dirs,
    claim_file,
    move_to_review,
    move_to_failed,
    move_to_archive,
    save_manifest,
    get_worker_id,
)

logger = structlog.get_logger(__name__)

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff", ".tif"}
PDF_EXTENSIONS = {".pdf"}
SUPPORTED_EXTENSIONS = IMAGE_EXTENSIONS | PDF_EXTENSIONS | {".txt", ".md", ".docx"}


def _build_checkpointer():
    """MemorySaver by default; a process-level compiled graph (see
    ``get_compiled_graph``) keeps that saver alive so ``interrupt()`` HITL
    can ``Command(resume=...)`` in the same process (API embeds the watcher).

    The filesystem review bin remains the durable park across process
    restart: when MemorySaver is empty, ``resume_from_review`` falls back to
    a fresh extract invoke from the manifest.

    Set MAILROOM_CHECKPOINTER=sqlite to opt into the on-disk checkpointer
    (debugging / resume-across-restart). Falls back to MemorySaver if SQLite
    is unavailable.
    """
    import os

    if os.environ.get("MAILROOM_CHECKPOINTER", "memory") != "sqlite":
        logger.info("checkpointer_initialized", backend="memory")
        return MemorySaver()
    try:
        import sqlite3
        from langgraph.checkpoint.sqlite import SqliteSaver
        from pipeline.bins import get_base_dir

        db_path = get_base_dir() / "checkpoints.db"
        db_path.parent.mkdir(parents=True, exist_ok=True)
        conn = sqlite3.connect(str(db_path), check_same_thread=False)
        checkpointer = SqliteSaver(conn)
        try:
            checkpointer.setup()
        except Exception:
            pass
        logger.info("checkpointer_initialized", backend="sqlite", path=str(db_path))
        return checkpointer
    except Exception:
        logger.warning("sqlite_checkpointer_unavailable", fallback="memory")
    return MemorySaver()


_compiled_graph = None
_compiled_graph_lock = threading.Lock()


def get_compiled_graph():
    """Process-level compiled graph + checkpointer.

    ``build_graph()`` still returns a fresh graph for tests that invoke it
    directly. Pipeline entrypoints (``run_pipeline`` / ``resume_from_review``)
    share this singleton so a MemorySaver checkpoint survives parking in
    ``review/`` until ``Command(resume=...)``.
    """
    global _compiled_graph
    with _compiled_graph_lock:
        if _compiled_graph is None:
            _compiled_graph = build_graph()
        return _compiled_graph


def reset_compiled_graph():
    """Drop the process-level graph (tests change MAILROOM_BASE_DIR / saver)."""
    global _compiled_graph
    with _compiled_graph_lock:
        _compiled_graph = None


def _ensure_dirs():
    ensure_dirs(
        inbox_dir(),
        processing_dir(),
        classified_dir(),
        review_dir(),
        failed_dir(),
        archive_dir(),
        manifests_dir(),
    )


def _read_file_text(file_path: Path) -> tuple[str, bool]:
    ext = file_path.suffix.lower()
    if ext in IMAGE_EXTENSIONS:
        return _extract_text_from_image(file_path)
    if ext in PDF_EXTENSIONS:
        return _extract_text_from_pdf(file_path)
    if ext == ".docx":
        return _extract_text_from_docx(file_path)
    try:
        text = file_path.read_text(errors="replace")
        if not text.strip():
            return ("", False)
        return (text, True)
    except Exception:
        try:
            text = file_path.read_bytes().decode("utf-8", errors="replace")
            return (text, bool(text.strip()))
        except Exception:
            return (f"[Unreadable file: {file_path.name}]", False)


def _render_doc_pages(file_path: Path) -> list[str]:
    """Render an input document to page-image data-URIs for vision-capable
    agents (PDFs page-by-page, image files passed through). Empty when vision is
    disabled or rendering is unavailable — text-only behaviour is unchanged."""
    try:
        from llm.vision import render_document_pages

        return render_document_pages(file_path)
    except Exception:
        logger.exception("doc_page_render_failed", file=str(file_path))
        return []


def _extract_text_from_docx(file_path: Path) -> tuple[str, bool]:
    """Extract text from .docx (paragraphs + tables) via python-docx.

    Previously .docx files fell through to the generic reader and were decoded
    as UTF-8 — i.e. zip binary garbage — which the classifier then tried to
    label. Unreadable files return the standard unreadable marker with
    ok=False so the pipeline routes them to review instead of misclassifying.
    """
    try:
        from docx import Document

        doc = Document(str(file_path))
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        if text.strip():
            logger.info("docx_text_extracted", file=file_path.name, chars=len(text))
            return (text, True)
        return ("", False)
    except ImportError:
        logger.warning("python_docx_missing", file=str(file_path))
        return (f"[Unreadable file: {file_path.name} — .docx support requires python-docx]", False)
    except Exception:
        logger.exception("docx_extraction_failed", file=str(file_path))
        return (f"[Unreadable file: {file_path.name} — not a valid .docx]", False)


def _extract_text_from_image(file_path: Path) -> tuple[str, bool]:
    logger.info("image_detected", file=str(file_path))
    try:
        from agents.image_extractor import ImageExtractor
        from observability.tracing import observation

        extractor = ImageExtractor()
        with observation("extract-image-text", as_type="retriever", input={"file": file_path.name}) as span:
            result = extractor.extract(file_path)
            if span is not None:
                span.update(output={"chars": len(result.get("text", ""))})
        text = result.get("text", "")
        if text:
            logger.info("image_extracted", file=file_path.name, chars=len(text))
            return (text, True)
    except Exception:
        logger.exception("image_extraction_failed", file=str(file_path))
    return (f"[Image file: {file_path.name} — text extraction failed]", False)


def _extract_text_from_pdf(file_path: Path) -> tuple[str, bool]:
    logger.info("pdf_detected", file=str(file_path))
    try:
        from agents.pdf_transcriber import PDFTranscriber
        from observability.tracing import observation

        transcriber = PDFTranscriber()
        with observation("transcribe-pdf", as_type="retriever", input={"file": file_path.name}) as span:
            result = transcriber.transcribe(file_path)
            if span is not None:
                span.update(
                    output={
                        "chars": len(result.get("markdown", "") or result.get("text", "")),
                        "method": result.get("method"),
                        "confidence": result.get("confidence"),
                    }
                )
        text = result.get("markdown", "") or result.get("text", "")
        if text:
            logger.info("pdf_transcribed", file=file_path.name, chars=len(text))
            return (text, True)
    except Exception:
        logger.exception("pdf_transcription_failed", file=str(file_path))
    try:
        import subprocess, tempfile
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            pass
        subprocess.run(["pdftotext", str(file_path), tmp.name], capture_output=True, timeout=30)
        text = Path(tmp.name).read_text(errors="replace")
        Path(tmp.name).unlink(missing_ok=True)
        if text.strip():
            logger.info("pdf_fallback_text", chars=len(text))
            return (text, True)
    except Exception:
        logger.exception("pdf_fallback_failed")
    return (f"[PDF file: {file_path.name} — transcription failed]", False)


def entry_route(state: dict) -> str:
    """Entry router: a review-resume re-invocation starts at fresh extraction
    (doc_type already known from the manifest); everything else goes through
    normal intake → classify.

    The `review_decision == "approved"` guard is deliberate: only the
    resume-from-review path sets it, so a crashed/partial run can never be
    mistaken for a resume and skip classification (pilot: correspondence_01
    ended with output=null when a degraded second run took the extract branch
    without a real classification).
    """
    if (
        state.get("resume_extraction")
        and state.get("review_decision") == "approved"
        and state.get("doc_type")
    ):
        return "extract"
    return "intake"


def _build_handoff_context(state: DocumentState) -> str | None:
    """Chained-eval handoff: prefix the sorter's classification (doc class +
    contract subtype) to the specialist's extraction call so it extracts with
    the expected field/clause set in mind (mirrors the sister repo's
    run_chained_eval pattern).

    KANBAN-063: when this extraction is an ARBITER-APPROVED RETRY, the
    arbiter's field fix-list rides along too, so the specialist repairs the
    named failures instead of rolling the dice on a blind re-run."""
    doc_type = state.get("doc_type")
    if not doc_type:
        return None
    context = f"Sorter classification: doc_type={doc_type}"
    extract_class = doc_type
    try:
        from pipeline.config import resolve_extract_class

        resolved = resolve_extract_class(doc_type)
        if resolved:
            extract_class = resolved
        if extract_class and extract_class != doc_type:
            context += f" extract_class={extract_class}"
    except Exception:
        logger.debug("extract_class_resolve_failed", doc_type=doc_type, exc_info=True)
    contract_subtype = state.get("contract_subtype")
    doc_subclass = state.get("doc_subclass")
    subtype = contract_subtype or doc_subclass
    if (doc_type == "contract" or extract_class == "contract") and contract_subtype:
        context += f" contract_subtype={contract_subtype}"
        context += (
            " CUAD family extraction: capture that family's characteristic "
            "operative clauses as present-only cuad_clauses lines "
            "'<Category>: <short evidence>' (grant of license, resale/purchase, "
            "franchise fees, maintenance/support, joint-venture sharing, "
            "non-compete covenants, etc. as the family requires)."
        )
    elif doc_subclass:
        context += f" doc_subclass={doc_subclass}"
    if doc_type == "merger_agreement":
        context += (
            " MAUD extraction: set merger_consideration / contract_value to the "
            "merger-consideration token all_cash|all_stock|mixed_cash_stock|"
            "mixed_cash_stock_election|other; put answered MAUD questions into "
            "maud_clauses as '<Question>: <short evidence>'."
        )
    confidence = state.get("classification_confidence")
    if confidence is not None:
        context += f" confidence={float(confidence):.2f}"
    # Relations clerk (HUB-040): the archive's advisory RELATED block — what
    # this document/matter is already known to relate to (zero-shot lift).
    try:
        from pipeline.relations import context_block

        related = context_block(
            matter_id=state.get("matter_id"), doc_id=state.get("doc_id")
        )
        if related:
            context += "\n" + related
    except Exception:
        logger.debug("relations_handoff_context_failed")
    if state.get("arbiter_retry_count"):
        findings = list(state.get("judge_findings") or [])
        to_fix = [str(f) for f in (state.get("arbiter_fields_to_fix") or []) if f]
        if to_fix:
            findings = [f"fields_to_fix: {', '.join(to_fix)}"] + findings
        handoff = state.get("arbiter_handoff")
        if handoff:
            findings.append(str(handoff)[:400])
        context += (
            "\nARBITER RETRY — the previous extraction was rejected by quality "
            "review. Fix these specific problems: "
            + ("; ".join(str(f) for f in findings)[:1200] or "(unspecified)")
            + "."
        )
    context += ". Extract this document's fields accordingly, ensuring every expected item of this document class is captured."
    try:
        from langchain_agents.doc_inventories import specialist_handoff

        extra = specialist_handoff(doc_type, subtype)
        if extra:
            context += " " + extra
    except Exception:
        pass
    return context


def _enrich_contract_result(result: dict | None, state: dict) -> dict:
    """Attach Hub inventory fields after the specialist returns."""
    payload = dict(result or {})
    doc_type = state.get("doc_type") or ""
    extract_class = doc_type
    try:
        from pipeline.config import resolve_extract_class

        extract_class = resolve_extract_class(doc_type) or doc_type
    except Exception:
        pass
    try:
        from langchain_agents.doc_inventories import enrich_extraction

        return enrich_extraction(
            payload,
            doc_type=doc_type,
            extract_class=extract_class,
            subtype=state.get("contract_subtype") or state.get("doc_subclass"),
        )
    except Exception:
        logger.exception("contract_inventory_enrich_failed")
        return payload


def _specialist_extractor_map():
    """Name → extract function. Keys MUST match taxonomy ``specialist:`` values."""
    return {
        "contracts_specialist": _extract_contracts,
        "corporate_records_specialist": _extract_corporate_records,
        "correspondence_specialist": _extract_correspondence,
        "compliance_specialist": _extract_compliance,
        "insurance_claims_specialist": _extract_insurance_claims,
    }


def _build_specialist_dispatch():
    from pipeline.config import load_config

    cfg = load_config()
    doc_classes = cfg.get("doc_classes", [])
    extractors = _specialist_extractor_map()
    dispatch = {}
    unmapped = []
    for cls in doc_classes:
        key = cls["key"]
        spec_name = cls.get("specialist", "")
        fn = extractors.get(spec_name)
        if fn is None:
            unmapped.append((key, spec_name))
            continue
        dispatch[key] = fn
    expected = {c["key"] for c in doc_classes}
    if unmapped or set(dispatch) != expected:
        raise RuntimeError(
            "specialist dispatch incomplete for taxonomy classes: "
            f"unmapped={unmapped} missing={sorted(expected - set(dispatch))}"
        )
    return dispatch


def _specialist_memory_name(doc_type: str) -> str | None:
    """Resolve the specialist agent name for per-agent memory.

    Unmapped / retired / unknown types return None — never fall back to
    contracts_specialist (that attributed the wrong agent's outcomes).
    Extract aliases keep ``state['doc_type']``. ``merger_agreement`` is a
    live taxonomy key and dispatches through its own specialist mapping
    (shared ``contracts_specialist``).
    """
    try:
        from pipeline.config import load_config, resolve_extract_class

        resolved = resolve_extract_class(doc_type) or doc_type
        for cls in load_config().get("doc_classes", []) or []:
            if cls.get("key") == resolved and cls.get("specialist"):
                return str(cls["specialist"])
    except Exception:
        pass
    return None


def _extract_dispatch_key(doc_type: str) -> str:
    """Taxonomy dispatch key for extract — aliases keep ``state['doc_type']``."""
    from pipeline.config import resolve_extract_class

    return resolve_extract_class(doc_type) or doc_type


def _unsupported_extraction_update(state: DocumentState, doc_type: str) -> dict[str, Any]:
    """Park a missing-specialist extract for human review (no retry)."""
    retry_max = get_confidence_thresholds().get("retry_max", 1)
    attempts = max(state.get("extraction_attempts", 0) or 0, retry_max) + 1
    reason = f"no specialist dispatched for doc_type={doc_type!r}"
    logger.error("no_specialist_dispatch", doc_type=doc_type, doc_id=state.get("doc_id"))
    return {
        "extracted_data": {"_unsupported": True},
        "extraction_confidence": 0.0,
        "extraction_attempts": attempts,
        "extraction_guardrail": ["no_specialist_dispatch"],
        "conflict_detected": False,
        "conflict_details": [],
        "escalation_reason": reason,
        "transient_error": False,
    }


def _chunk_config() -> dict:
    """Chunked-extraction config from taxonomy.yaml (`chunking:` block).

    Chunking (v15+ vendored architecture) splits documents longer than
    ``chunk_chars`` into overlapping windows, extracts each window, and merges
    deterministically — nothing is truncated. Documents that fit in a single
    window take the plain single-pass path, so small documents are unaffected.
    """
    try:
        from pipeline.config import load_config

        return load_config().get("chunking", {}) or {}
    except Exception:
        return {}


def _specialist_input_budget(agent) -> int:
    """Chars the specialist will actually send — cap chunk windows to this."""
    if hasattr(agent, "_configured_max_input_chars"):
        try:
            return int(agent._configured_max_input_chars())
        except Exception:
            pass
    budget = getattr(agent, "_max_input_chars", None)
    if budget:
        return int(budget)
    return 90_000


def _instantiate_specialist(agent_fn, handoff_context):
    """Construct a specialist; pipeline agents do not take handoff_context."""
    try:
        return agent_fn(handoff_context=handoff_context)
    except TypeError:
        agent = agent_fn()
        if handoff_context is not None:
            agent.handoff_context = handoff_context
        return agent


def _run_chunked_extraction(agent_fn, doc_text, pages, handoff_context):
    """Run a specialist extraction, chunking long documents (v15+ pass).

    ``extract_chunked`` falls through to the plain single-pass ``extract``
    when the document fits in one window, so this never changes small-document
    output. ``pages`` (MAILROOM PATCH) are attached to the first chunk only.
    Window size is capped at the agent's ``max_input_chars`` so a chunk is
    never silently truncated inside ``extract()``.
    """
    cfg = _chunk_config()
    chunk_chars = int(cfg.get("chunk_chars", 90_000))
    overlap_chars = int(cfg.get("overlap_chars", 8_000))
    agent = _instantiate_specialist(agent_fn, handoff_context)
    budget = _specialist_input_budget(agent)
    overlap_chars = min(overlap_chars, max(0, budget // 8))
    chunk_chars = min(chunk_chars, max(1_000, budget - overlap_chars))
    if cfg.get("enabled", True) and hasattr(agent, "extract_chunked"):
        kwargs = dict(
            doc_text=doc_text,
            chunk_chars=chunk_chars,
            overlap_chars=overlap_chars,
            pages=pages,
        )
        try:
            return agent.extract_chunked(**kwargs, handoff_context=handoff_context)
        except TypeError:
            return agent.extract_chunked(**kwargs)
    try:
        return agent.extract(doc_text, pages=pages, handoff_context=handoff_context)
    except TypeError:
        return agent.extract(doc_text, pages=pages)


def intake_node(state: DocumentState) -> dict[str, Any]:
    """The intake node — the FIRST node of the pipeline (HUB-038).

    The INTAKE agent IS the ingest specialist: "ingest" and "intake" are the
    same step (the split was an unintentional naming mistake — unified
    2026-09-03). This one node performs the full ingest + intake work:

    1. Claim + transcribe the file (``_read_file_text``).
    2. Run the deterministic intake clerk (``apply_intake`` — the dojo gold
       baseline, never skipped).
    3. Run the LLM-assisted intake agent (``IntakeAgent``) when gated
       (messy / over-sorter-budget): TRIAGE (advisory read), CLEAN
       (structural repair, re-normalized), PREPARE (section map) — sliding
       windows, never truncated.
    4. Write the processing manifest + catalog record + the ``ingested``
       compliance audit event (A-1/A-7).

    The intake work product is what the sorter depends on: the advisory
    triage rides ``state.intake_prep`` and is fed to ``classify`` as a labeled
    prior; cleaning and section maps refine the text the sorter reads.
    """
    _ensure_dirs()
    worker_id = get_worker_id()

    if state.get("file_path"):
        file_path = Path(state["file_path"])
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
    else:
        inbox = inbox_dir()
        files = list(inbox.glob("*"))
        if not files:
            raise RuntimeError("No files in inbox")
        file_path = files[0]

    try:
        in_inbox = file_path.resolve().is_relative_to(inbox_dir().resolve())
    except (OSError, ValueError):
        in_inbox = False
    if in_inbox:
        file_path = claim_file(file_path, worker_id)

    doc_text, text_ok = _read_file_text(file_path)
    from agents.intake import apply_intake, llm_intake_enabled, should_llm_intake

    raw_text = doc_text
    doc_text, intake_stats = apply_intake(doc_text, filename=file_path.name)
    # HUB-038: LLM-assisted intake (triage + clean + prepare) — gated to
    # messy / over-sorter-budget documents; sliding windows, NEVER truncates.
    intake_prep = None
    try:
        if should_llm_intake(doc_text, intake_stats) and llm_intake_enabled():
            from agents.intake import IntakeAgent
            from observability.tracing import observation

            intake_agent = IntakeAgent()
            with observation(
                "intake-llm-prep",
                as_type="span",
                input={"file": file_path.name, "chars": len(doc_text)},
            ) as span:
                intake_prep = intake_agent.intake_run(doc_text, filename=file_path.name)
                if span is not None:
                    triage = intake_prep.get("triage") or {}
                    span.update(
                        output={
                            "triage_class": triage.get("primary_doc_class"),
                            "triage_confidence": triage.get("confidence"),
                            "sections": len(intake_prep.get("sections") or []),
                            "windows": intake_prep.get("windows", 1),
                            "cleaned": bool(intake_prep.get("cleaned")),
                            "changed": bool(intake_prep.get("changed")),
                        }
                    )
            if intake_prep.get("cleaned"):
                from llm_dojo_scoring.intake import looks_messy as _looks_messy

                doc_text = intake_prep["cleaned"]
                intake_stats = dict(intake_prep["clean_stats"])
                intake_stats["messy"] = _looks_messy(doc_text, intake_stats)
                intake_stats["method"] = "llm"
    except Exception:
        logger.exception("intake_llm_failed", file=file_path.name)
        intake_prep = None
    try:
        from observability.suite_scoring import score_and_log_intake

        score_and_log_intake(raw_text, doc_text, intake_stats)
    except Exception:
        logger.debug("intake_suite_score_failed", exc_info=True)
    doc_pages = _render_doc_pages(file_path)

    matter_id = state.get("matter_id", "DEFAULT")
    intake_meta = state.get("intake_meta") or None
    if intake_meta is not None and intake_prep:
        intake_meta = dict(intake_meta)
        if intake_prep.get("triage"):
            intake_meta["triage"] = intake_prep["triage"]
        if intake_prep.get("sections") is not None:
            intake_meta["prep"] = {
                "section_count": len(intake_prep.get("sections") or []),
                "roles": sorted({s["role"] for s in intake_prep.get("sections") or []}),
                "windows": intake_prep.get("windows", 1),
            }
    manifest = DocumentManifest(
        matter_id=matter_id,
        original_filename=file_path.name,
        stage=PipelineStage.PROCESSING,
        trace_id=state.get("trace_id"),
        intake=intake_meta or None,
    )
    manifest.touch()
    save_manifest(manifest)

    logger.info(
        "intake",
        doc_id=manifest.doc_id,
        file=file_path.name,
        chars=len(doc_text),
        suffix=file_path.suffix,
        vision_pages=len(doc_pages),
        intake_changed=intake_stats.get("changed"),
        intake_messy=intake_stats.get("messy"),
    )
    # Write the processing-stage catalog record immediately so a crashed run is
    # visible to stuck-doc detection (`get_stuck_documents`) and `/ops/status`
    # instead of disappearing from the conveyor entirely.
    _catalog_upsert(
        {
            "doc_id": manifest.doc_id,
            "matter_id": matter_id,
            "original_filename": file_path.name,
            "doc_type": None,
            "stage": PipelineStage.PROCESSING.value,
        },
        stage=PipelineStage.PROCESSING.value,
    )
    # A-1/A-7: ingest is the start of the compliance record — record the file
    # hash (sha256) + size so post-archive substitution is detectable.
    file_hash = _file_sha256(file_path)
    ingest_detail = {"file_sha256": file_hash, "size_bytes": _file_size(file_path)}
    _emit_stage_audit(
        {"doc_id": manifest.doc_id, "matter_id": matter_id, "stage": PipelineStage.PROCESSING.value,
         "original_filename": file_path.name},
        "ingested",
        actor="pipeline",
        detail=ingest_detail,
    )
    # Carry the hash in state so archive verification + provenance persistence
    # can compare against it (A-7).
    intake_state = {
        "doc_id": manifest.doc_id,
        "matter_id": matter_id,
        "original_filename": file_path.name,
        "stage": PipelineStage.PROCESSING.value,
        "file_path": str(file_path),
        "doc_text": doc_text,
        "doc_pages": doc_pages,
        "intake_messy": bool(intake_stats.get("messy")),
        "intake_changed": bool(intake_stats.get("changed")),
        "intake_prep": intake_prep,
        "classification_attempts": 0,
        "extraction_attempts": 0,
        "retry_count": 0,
        "conflict_detected": False,
        "error_message": None if text_ok else f"Could not extract text from {file_path.suffix} file",
        "file_sha256": file_hash,
        "size_bytes": _file_size(file_path),
    }
    _catalog_upsert(
        {
            "doc_id": manifest.doc_id,
            "matter_id": matter_id,
            "original_filename": file_path.name,
            "doc_type": None,
            "stage": PipelineStage.PROCESSING.value,
            "file_sha256": file_hash,
        },
        stage=PipelineStage.PROCESSING.value,
    )
    return intake_state


def classify_node(state: DocumentState) -> dict[str, Any]:
    doc_text = state.get("doc_text", "")
    if not doc_text or not doc_text.strip():
        logger.warning("empty_doc_text_classify", doc_id=state.get("doc_id"))
        # Empty/unreadable text can never be classified — route straight to
        # human review instead of burning a retry call on the same empty text
        # (which also clobbered this escalation reason on the retry). Setting
        # classification_attempts past retry_max makes after_classify send it
        # to review immediately.
        retry_max = get_confidence_thresholds().get("retry_max", 1)
        return {
            "doc_type": UNKNOWN_DOC_TYPE,
            "classification_confidence": 0.1,
            "classification_attempts": max(state.get("classification_attempts", 0), retry_max) + 1,
            "stage": PipelineStage.CLASSIFIED.value,
            "escalation_reason": "Empty or unreadable document content",
            "transient_error": False,
        }

    from agents.sorter import SorterAgent
    from llm.retry import is_transient_error
    from pipeline.guards import apply_classification_guard

    sorter = SorterAgent()
    attempts = state.get("classification_attempts", 0)
    try:
        # HUB-038: the advisory intake read rides as a labeled prior — the
        # sorter verifies independently (the vendored sorter_v14 prompt is
        # never mutated). Over-budget documents slide through windows inside
        # the sorter subclass — never truncated.
        from agents.intake import format_intake_prior

        intake_prior = format_intake_prior(state.get("intake_prep") or None)
        # Structured classify includes per-class doc_subclass (dojo catalogs).
        classified = sorter.classify_json(
            doc_text,
            pages=state.get("doc_pages"),
            intake_prior=intake_prior,
        )
        doc_type = classified.get("doc_type") or ""
        contract_subtype = classified.get("contract_subtype")
        doc_subclass = classified.get("doc_subclass")
        try:
            confidence = float(classified.get("confidence", 0.5))
        except (TypeError, ValueError):
            confidence = 0.5
        reasoning = classified.get("reasoning") or ""
    except Exception as exc:
        if is_transient_error(exc):
            # Provider-side transient failure (connection/timeout/rate-limit/
            # 5xx). Do NOT increment the confidence retry budget; routing
            # retries this same node via the `classify` self-loop.
            transient = state.get("transient_retries_classify", 0) + 1
            logger.warning(
                "classify_transient_error",
                doc_id=state.get("doc_id"),
                error=str(exc)[:300],
                transient_retries=transient,
            )
            return {
                "transient_error": True,
                "transient_retries_classify": transient,
                "classification_attempts": attempts,
                "stage": PipelineStage.CLASSIFIED.value,
                "error_message": f"transient provider error: {str(exc)[:200]}",
                "escalation_reason": "transient provider error during classification",
            }
        # L-15: non-transient classify errors used to `raise` — the run's
        # catch-all sent the document to the FAILED bin. Classify hard-failures
        # are exactly the documents needing human eyes, so mirror extract's
        # conversion: surface the error as a low-confidence review-routing
        # result instead of crashing the run. The classification_attempts
        # counter is pushed past retry_max so after_classify routes straight to
        # human review (same mechanism as the empty-text fast path above).
        retry_max = get_confidence_thresholds().get("retry_max", 1)
        logger.exception(
            "classification_exception",
            doc_id=state.get("doc_id"),
            error=str(exc)[:300],
        )
        return {
            "doc_type": UNKNOWN_DOC_TYPE,
            "contract_subtype": None,
            "doc_subclass": None,
            "classification_confidence": 0.1,
            "classification_attempts": max(attempts, retry_max) + 1,
            "stage": PipelineStage.CLASSIFIED.value,
            "error_message": f"classification error: {str(exc)[:200]}",
            "escalation_reason": f"classification failed ({type(exc).__name__}) — routing to human review",
            "transient_error": False,
        }
    attempts = attempts + 1

    guard, confidence = apply_classification_guard(
        {
            "doc_type": doc_type,
            "classification_confidence": confidence,
            "contract_subtype": contract_subtype,
            "doc_subclass": doc_subclass,
        }
    )
    if not guard["ok"]:
        # Confidence is already clamped by apply_classification_guard. Leave
        # doc_type untouched so routing's unknown-type check can still send it
        # to human review. Record the guardrail on state so it is scored
        # (guardrail_triggered) and visible in traces.
        # Per-agent memory (iterative improvement): record what the guardrail
        # rejected so retry calls can learn from it.
        try:
            from langchain_agents.memory import record_outcome

            record_outcome(
                "sorter",
                doc_type=doc_type or "",
                decision=f"{doc_type}/{doc_subclass or contract_subtype}",
                confidence=confidence,
                feedback=f"classification guardrail rejected: {guard['issues']}",
                source="guardrail",
            )
        except Exception:
            pass

    logger.info(
        "classified",
        doc_type=doc_type,
        contract_subtype=contract_subtype,
        doc_subclass=doc_subclass,
        confidence=confidence,
        attempts=attempts,
    )
    result = {
        "doc_type": doc_type,
        "contract_subtype": contract_subtype,
        "doc_subclass": doc_subclass,
        "classification_confidence": confidence,
        "classification_attempts": attempts,
        "classification_guardrail": guard["issues"],
        "stage": PipelineStage.CLASSIFIED.value,
        "escalation_reason": reasoning
        if confidence < get_confidence_thresholds().get("high", 0.95)
        else None,
        "transient_error": False,
    }
    # A-1: the classification decision (incl. guardrail issues) is part of the
    # compliance record.
    _emit_stage_audit({**state, **result}, "classified", actor="sorter",
                      detail={"attempts": attempts, "guardrail_issues": guard["issues"] or None})
    return result


def retry_classify_node(state: DocumentState) -> dict[str, Any]:
    from agents.sorter import SorterAgent
    from llm.retry import is_transient_error

    sorter = SorterAgent()
    doc_text = state.get("doc_text", "")
    attempts = state.get("classification_attempts", 0)

    prev_type = state.get("doc_type", "")
    prev_confidence = state.get("classification_confidence", 0)
    # Inner-context memory (iterative improvement): the retry sees what similar
    # classifications were corrected/guarded before.
    try:
        from langchain_agents.memory import recent_context

        memory = recent_context("sorter", doc_type=prev_type or "", k=3)
    except Exception:
        memory = ""
    preamble = (
        f"RE-EVALUATION REQUESTED - previous classification was '{prev_type}' with "
        f"confidence {prev_confidence:.2f}. Please re-examine this document independently:"
    )
    if memory:
        preamble = f"{preamble}\n\n{memory}"
    from agents.intake import format_intake_prior

    try:
        # HUB-038: no truncation — the retry reads the FULL document through
        # sliding windows (preamble + advisory intake prior on every window).
        classified = sorter.classify_json(
            doc_text,
            pages=state.get("doc_pages"),
            prefix=preamble,
            intake_prior=format_intake_prior(state.get("intake_prep") or None),
        )
        doc_type = classified.get("doc_type") or ""
        contract_subtype = classified.get("contract_subtype")
        doc_subclass = classified.get("doc_subclass")
        try:
            confidence = float(classified.get("confidence", 0.5))
        except (TypeError, ValueError):
            confidence = 0.5
        reasoning = classified.get("reasoning") or ""
    except Exception as exc:
        if is_transient_error(exc):
            transient = state.get("transient_retries_retry_classify", 0) + 1
            logger.warning(
                "retry_classify_transient_error",
                doc_id=state.get("doc_id"),
                error=str(exc)[:300],
                transient_retries=transient,
            )
            return {
                "transient_error": True,
                "transient_retries_retry_classify": transient,
                "classification_attempts": attempts,
                "stage": PipelineStage.CLASSIFIED.value,
                "error_message": f"transient provider error: {str(exc)[:200]}",
                "escalation_reason": "transient provider error during re-classification",
            }
        # L-15: match classify_node — a hard failure during re-classification
        # must park for human review, not crash the run into the failed bin.
        retry_max = get_confidence_thresholds().get("retry_max", 1)
        logger.exception(
            "retry_classification_exception",
            doc_id=state.get("doc_id"),
            error=str(exc)[:300],
        )
        return {
            "classification_confidence": 0.1,
            "classification_attempts": max(attempts, retry_max) + 1,
            "stage": PipelineStage.CLASSIFIED.value,
            "error_message": f"re-classification error: {str(exc)[:200]}",
            "escalation_reason": (
                f"re-classification failed ({type(exc).__name__}) — routing to human review"
            ),
            "transient_error": False,
        }
    attempts = attempts + 1

    from pipeline.guards import apply_classification_guard

    guard, confidence = apply_classification_guard(
        {
            "doc_type": doc_type,
            "classification_confidence": confidence,
            "contract_subtype": contract_subtype,
            "doc_subclass": doc_subclass,
        }
    )

    logger.info(
        "retry_classified",
        doc_type=doc_type,
        contract_subtype=contract_subtype,
        doc_subclass=doc_subclass,
        confidence=confidence,
        attempts=attempts,
    )
    return {
        "doc_type": doc_type,
        "contract_subtype": contract_subtype,
        "doc_subclass": doc_subclass,
        "classification_confidence": confidence,
        "classification_attempts": attempts,
        "retry_count": state.get("retry_count", 0) + 1,
        "classification_guardrail": guard["issues"],
        "stage": PipelineStage.CLASSIFIED.value,
        "escalation_reason": reasoning
        if confidence < get_confidence_thresholds().get("high", 0.95)
        else None,
        "transient_error": False,
    }


def review_classify_node(state: DocumentState) -> dict[str, Any]:
    """KANBAN-062 (Lane A): independent agent second opinion on a medium-band
    classification.

    The reviewer classifies the document BLIND (no hint of the sorter's
    answer — independence is the point). Agreement/override is computed here
    in code. The lane never changes the failure surface: reviewer high-
    confidence → extract with the reviewer's label applied; anything else →
    human_review with BOTH opinions preserved on state.
    """
    from agents.sorter_reviewer import SorterReviewerAgent
    from llm.retry import is_transient_error
    from pipeline.config import get_sorter_label_set
    from pipeline.guards import apply_classification_guard
    from langchain_agents.sorter_agent import CONTRACT_SUBTYPE_KEYS, SUBTYPE_UNKNOWN

    doc_text = state.get("doc_text", "")
    sorter_type = state.get("doc_type")
    sorter_confidence = state.get("classification_confidence")

    try:
        reviewer = SorterReviewerAgent()
        result = reviewer.review(
            doc_text,
            pages=state.get("doc_pages"),
            valid_doc_types=sorted(get_sorter_label_set()),
            contract_subtypes=list(CONTRACT_SUBTYPE_KEYS) + [SUBTYPE_UNKNOWN],
        )
    except Exception as exc:
        if is_transient_error(exc):
            transient = state.get("transient_retries_review_classify", 0) + 1
            logger.warning(
                "review_classify_transient_error",
                doc_id=state.get("doc_id"),
                error=str(exc)[:300],
                transient_retries=transient,
            )
            return {
                "transient_error": True,
                "transient_retries_review_classify": transient,
                "stage": PipelineStage.CLASSIFIED.value,
                "error_message": f"transient provider error: {str(exc)[:200]}",
                "escalation_reason": "transient provider error during sorter review",
            }
        # Reviewer hard-failed: escalate with the sorter's original answer
        # intact (fail-safe — same destination the doc had before this lane).
        # HUB-043: the exception TEXT rides the reason so the completion echo
        # can translate the actual cause for the recipient (an opaque
        # "(RuntimeError)" alone told the sender nothing).
        logger.exception("review_classify_failed", doc_id=state.get("doc_id"))
        return {
            "review_verdict": "reviewer_error",
            "stage": PipelineStage.CLASSIFIED.value,
            "escalation_reason": f"sorter reviewer failed ({type(exc).__name__}: {str(exc)[:160]}) — routing to human review",
            "transient_error": False,
        }

    reviewer_type = result.get("doc_type")
    reviewer_subtype = result.get("contract_subtype")
    reviewer_subclass = result.get("doc_subclass")
    try:
        reviewer_confidence = float(result.get("confidence", 0.0))
    except (TypeError, ValueError):
        reviewer_confidence = 0.0
    # Same deterministic guard the sorter's answers pass through: never trust
    # an out-of-range confidence or unknown class from the reviewer either.
    # Clamping happens here so a high-confidence invalid subtype cannot win
    # the lane and auto-extract.
    _guard, reviewer_confidence = apply_classification_guard(
        {
            "doc_type": reviewer_type,
            "classification_confidence": reviewer_confidence,
            "contract_subtype": reviewer_subtype,
            "doc_subclass": reviewer_subclass,
        }
    )
    reviewer_reasoning = str(result.get("reasoning", ""))

    # Class-aware high: agree uses sorter class; override uses the reviewer's
    # proposed class (severity of the label they want to win).
    agree_high = get_confidence_thresholds(sorter_type).get("high", 0.97)
    override_high = get_confidence_thresholds(reviewer_type).get("high", 0.97)
    if reviewer_type == sorter_type:
        verdict = (
            "reviewer_agrees_high"
            if reviewer_confidence >= agree_high
            else "reviewer_agrees_low"
        )
    else:
        verdict = (
            "reviewer_overrides"
            if reviewer_confidence >= override_high
            else "reviewer_conflicts"
        )

    logger.info(
        "review_classified",
        doc_id=state.get("doc_id"),
        sorter_type=sorter_type,
        sorter_confidence=sorter_confidence,
        reviewer_type=reviewer_type,
        reviewer_confidence=reviewer_confidence,
        verdict=verdict,
    )
    updates = {
        "reviewer_doc_type": reviewer_type,
        "reviewer_contract_subtype": reviewer_subtype,
        "reviewer_doc_subclass": reviewer_subclass,
        "reviewer_confidence": reviewer_confidence,
        "review_verdict": verdict,
        # The reviewer's reasoning rides on escalation_reason so a human
        # reviewing the doc sees BOTH opinions in the manifest/catalog.
        "escalation_reason": (
            f"sorter review: sorter='{sorter_type}' "
            f"({float(sorter_confidence or 0):.2f}) → reviewer='{reviewer_type}' "
            f"({reviewer_confidence:.2f}, {verdict}): {reviewer_reasoning[:400]}"
        ),
        "stage": PipelineStage.CLASSIFIED.value,
        "transient_error": False,
    }
    # KANBAN-062: only a WINNING reviewer verdict may re-label the document.
    # reviewer_agrees_high re-asserts the (identical) sorter type with the
    # reviewer's confidence; reviewer_overrides replaces the sorter's label.
    # Low-confidence verdicts keep the sorter's answer untouched — the human
    # reviewer gets both opinions via the reviewer_* fields above.
    if verdict in ("reviewer_agrees_high", "reviewer_overrides"):
        updates["doc_type"] = reviewer_type
        updates["contract_subtype"] = reviewer_subtype
        updates["doc_subclass"] = reviewer_subclass
        updates["classification_confidence"] = reviewer_confidence
    return updates


def _fetch_matter_context(state: dict) -> list[dict]:
    """Best-effort fetch of archived matter records for the Boss / conflict
    detection. Returns a list of {doc_id, doc_type, extracted_data} dicts;
    never raises (DB unavailable → empty context)."""
    matter_id = state.get("matter_id")
    doc_id = state.get("doc_id")
    if not matter_id:
        return []
    try:
        from storage.catalog import get_matter_documents

        rows = _run_coro(lambda: get_matter_documents(matter_id))
        return [
            {
                "doc_id": r.doc_id,
                "doc_type": r.doc_type,
                "extracted_data": r.extracted_data or {},
                "stage": r.stage,
            }
            for r in rows
            if r.doc_id != doc_id and r.stage == "archived"
        ]
    except Exception:
        logger.exception("matter_context_fetch_failed")
        return []


def _normalized_compare(a, b) -> bool:
    """Normalized value comparison for conflict detection (list-aware)."""
    if isinstance(a, list) and isinstance(b, list):
        na = {_norm_str(x) for x in a}
        nb = {_norm_str(x) for x in b}
        return na == nb
    return _norm_str(a) == _norm_str(b)


def _norm_str(v) -> str:
    try:
        from observability.field_scoring import normalize_text

        return normalize_text(v)
    except Exception:
        return str(v).strip().lower()


def _detect_conflict(state: dict, extracted_data: dict | None) -> tuple[bool, list[str]]:
    """Deterministically compare a fresh extraction against archived records
    of the same matter. A conflict exists when the same field is populated on
    both sides with a different normalized value (e.g. two contracts in one
    matter claiming different governing laws or parties).

    Returns (conflict_detected, details). Best-effort: no DB → no conflict.
    """
    if not extracted_data:
        return False, []
    # Only fields present in the schema are conflict-relevant (ignore pipeline
    # metadata keys like `_report`).
    schema_fields = set()
    try:
        from schemas.documents import get_extraction_schema

        model = get_extraction_schema(state.get("doc_type") or "")
        if model is not None:
            schema_fields = set(model.model_fields.keys())
    except Exception:
        pass

    details: list[str] = []
    current_type = state.get("doc_type") or ""
    for record in _fetch_matter_context(state):
        # Same-class only: shared field names across schemas (e.g. both
        # contract and corporate_record have `effective_date`) are not
        # contradictions — they are different documents. The Boss exists to
        # catch two *contracts* in one matter claiming different governing
        # laws / parties, not a bylaws filing vs an MSA.
        if current_type and record.get("doc_type") and record.get("doc_type") != current_type:
            continue
        prior = record.get("extracted_data") or {}
        if not prior:
            continue
        for field in sorted(schema_fields):
            try:
                from langchain_agents.doc_inventories import skip_conflict_field

                if skip_conflict_field(field):
                    continue
            except Exception:
                pass
            new_val = extracted_data.get(field)
            old_val = prior.get(field)
            if new_val is None or old_val is None:
                continue
            if isinstance(new_val, str) and not new_val.strip():
                continue
            if isinstance(old_val, str) and not old_val.strip():
                continue
            if isinstance(new_val, (list, dict)) and not new_val:
                continue
            if isinstance(old_val, (list, dict)) and not old_val:
                continue
            if _normalized_compare(new_val, old_val):
                continue
            details.append(
                f"field '{field}' differs from archived record "
                f"{record.get('doc_id', '?')}: '{old_val}' vs '{new_val}'"
            )
            break  # one conflict per prior record is enough to escalate
    return bool(details), details


def extract_node(state: DocumentState) -> dict[str, Any]:
    doc_type = state.get("doc_type", "")
    doc_text = state.get("doc_text", "")
    doc_pages = state.get("doc_pages") or []

    dispatch = _build_specialist_dispatch()
    extractor = dispatch.get(_extract_dispatch_key(doc_type))
    if extractor is None:
        return _unsupported_extraction_update(state, doc_type)
    handoff_context = _build_handoff_context(state)
    attempts = state.get("extraction_attempts", 0)
    try:
        result = extractor(doc_text, doc_pages, handoff_context)
        result = _enrich_contract_result(result, state)
    except Exception as exc:
        from llm.retry import is_transient_error

        if is_transient_error(exc):
            # Transient provider failure: retry the same node without burning
            # the extraction retry budget (routed via the `extract` self-loop).
            transient = state.get("transient_retries_extract", 0) + 1
            logger.warning(
                "extract_transient_error",
                doc_id=state.get("doc_id"),
                error=str(exc)[:300],
                transient_retries=transient,
            )
            return {
                "transient_error": True,
                "transient_retries_extract": transient,
                "extraction_attempts": attempts,
                "extraction_confidence": 0.0,
                "extracted_data": None,
                "stage": PipelineStage.CLASSIFIED.value,
                "error_message": f"transient provider error: {str(exc)[:200]}",
                "escalation_reason": "transient provider error during extraction",
            }
        # Non-transient exception: convert it into a parse-level failure so the
        # deterministic guardrail clamps confidence and routing sends the doc
        # to retry → review instead of crashing the run silently.
        logger.exception(
            "extraction_exception",
            doc_id=state.get("doc_id"),
            doc_type=doc_type,
            error=str(exc)[:300],
        )
        result = {"_parse_error": True, "_exception": str(exc), "confidence": 0.0}
    confidence = result.pop("confidence", None)
    attempts = attempts + 1

    from pipeline.guards import apply_extraction_guard

    guard, confidence = apply_extraction_guard(doc_type, result, confidence, attempts=attempts)

    # Deterministic conflict detection against archived matter records: a
    # materially different value for the same schema field (governing law,
    # parties, effective date, amounts) is escalated to the Boss for
    # adjudication instead of silently overwriting matter history.
    conflict_detected, conflict_details = _detect_conflict(state, result)

    logger.info(
        "extracted",
        doc_type=doc_type,
        confidence=confidence,
        attempts=attempts,
        conflict_detected=conflict_detected,
    )
    result_dict = {
        "extracted_data": result,
        "extraction_confidence": confidence,
        "extraction_attempts": attempts,
        "extraction_guardrail": guard["issues"],
        "conflict_detected": conflict_detected,
        "conflict_details": conflict_details,
        "escalation_reason": "; ".join(conflict_details)
        if conflict_detected
        else state.get("escalation_reason"),
        "transient_error": False,
    }
    # Per-agent memory (iterative improvement): record guardrail rejections and
    # extraction confidence so the specialist's retry calls learn from them.
    try:
        from langchain_agents.memory import record_outcome

        mem_name = _specialist_memory_name(doc_type)
        if mem_name:
            record_outcome(
                mem_name,
                doc_type=doc_type or "",
                decision=f"conf={confidence:.2f}",
                confidence=confidence,
                feedback=f"extraction guardrail: {guard['issues']}" if guard["issues"] else f"extraction confidence {confidence:.2f}",
                source="guardrail" if guard["issues"] else "run",
                detail={"conflict_detected": bool(conflict_detected)},
            )
    except Exception:
        pass
    # A-1: record the extraction decision (attempt, guardrail issues, conflict).
    _emit_stage_audit(
        {**state, **result_dict},
        "extracted",
        actor=state.get("doc_type", "specialist"),
        detail={"attempts": attempts, "guardrail_issues": guard["issues"] or None,
                "conflict_detected": bool(conflict_detected)},
    )
    return result_dict


def _extract_contracts(
    doc_text: str, pages: list[str] | None = None, handoff_context: str | None = None
) -> dict:
    from agents.contracts_specialist import ContractsSpecialist
    return _run_chunked_extraction(ContractsSpecialist, doc_text, pages, handoff_context)


def _extract_corporate_records(
    doc_text: str, pages: list[str] | None = None, handoff_context: str | None = None
) -> dict:
    from agents.corporate_records_specialist import CorporateRecordsSpecialist
    return _run_chunked_extraction(CorporateRecordsSpecialist, doc_text, pages, handoff_context)


def _extract_correspondence(
    doc_text: str, pages: list[str] | None = None, handoff_context: str | None = None
) -> dict:
    from agents.correspondence_specialist import CorrespondenceSpecialist
    return _run_chunked_extraction(CorrespondenceSpecialist, doc_text, pages, handoff_context)


def _extract_compliance(
    doc_text: str, pages: list[str] | None = None, handoff_context: str | None = None
) -> dict:
    from agents.compliance_specialist import ComplianceSpecialist
    return _run_chunked_extraction(ComplianceSpecialist, doc_text, pages, handoff_context)


def _extract_insurance_claims(
    doc_text: str, pages: list[str] | None = None, handoff_context: str | None = None
) -> dict:
    from agents.insurance_claims_specialist import InsuranceClaimsSpecialist
    return _run_chunked_extraction(InsuranceClaimsSpecialist, doc_text, pages, handoff_context)


def retry_extract_node(state: DocumentState) -> dict[str, Any]:
    doc_type = state.get("doc_type", "")
    doc_text = state.get("doc_text", "")
    doc_pages = state.get("doc_pages") or []
    prev_extracted = state.get("extracted_data", {})
    attempts = state.get("extraction_attempts", 0)

    # Inner-context memory (iterative improvement): the specialist's retry sees
    # what similar extractions were guarded/corrected before.
    try:
        from langchain_agents.memory import recent_context

        mem_name = _specialist_memory_name(doc_type)
        memory = recent_context(mem_name, doc_type=doc_type or "", k=3) if mem_name else ""
    except Exception:
        memory = ""
    augmented_prefix = (
        "RE-EXTRACTION REQUESTED - previous extraction was low-confidence. "
        "Please re-examine this document independently. "
        f"Previous attempt found: {prev_extracted}"
    )
    handoff_parts = [_build_handoff_context(state), augmented_prefix]
    if memory:
        handoff_parts.append(memory)
    handoff_context = "\n\n".join(p for p in handoff_parts if p)

    dispatch = _build_specialist_dispatch()
    extractor = dispatch.get(_extract_dispatch_key(doc_type))
    if extractor is None:
        update = _unsupported_extraction_update(state, doc_type)
        update["retry_count"] = state.get("retry_count", 0) + 1
        return update
    try:
        result = extractor(doc_text, doc_pages, handoff_context)
        result = _enrich_contract_result(result, state)
    except Exception as exc:
        from llm.retry import is_transient_error

        if is_transient_error(exc):
            transient = state.get("transient_retries_retry_extract", 0) + 1
            logger.warning(
                "retry_extract_transient_error",
                doc_id=state.get("doc_id"),
                error=str(exc)[:300],
                transient_retries=transient,
            )
            return {
                "transient_error": True,
                "transient_retries_retry_extract": transient,
                "extraction_attempts": attempts,
                "extraction_confidence": 0.0,
                "extracted_data": None,
                "stage": PipelineStage.CLASSIFIED.value,
                "error_message": f"transient provider error: {str(exc)[:200]}",
                "escalation_reason": "transient provider error during re-extraction",
            }
        logger.exception(
            "retry_extraction_exception",
            doc_id=state.get("doc_id"),
            doc_type=doc_type,
            error=str(exc)[:300],
        )
        result = {"_parse_error": True, "_exception": str(exc), "confidence": 0.0}
    confidence = result.pop("confidence", None)
    attempts = attempts + 1

    from pipeline.guards import apply_extraction_guard

    guard, confidence = apply_extraction_guard(doc_type, result, confidence, attempts=attempts)

    conflict_detected, conflict_details = _detect_conflict(state, result)

    logger.info(
        "retry_extracted",
        doc_type=doc_type,
        confidence=confidence,
        attempts=attempts,
        conflict_detected=conflict_detected,
    )
    return {
        "extracted_data": result,
        "extraction_confidence": confidence,
        "extraction_attempts": attempts,
        "retry_count": state.get("retry_count", 0) + 1,
        "extraction_guardrail": guard["issues"],
        "conflict_detected": conflict_detected,
        "conflict_details": conflict_details,
        "escalation_reason": "; ".join(conflict_details)
        if conflict_detected
        else state.get("escalation_reason"),
        "transient_error": False,
    }


def _clean_fields_for_judge(extracted: dict | None) -> dict:
    """Drop pipeline metadata keys before judge/arbiter input (mirrors the
    `_emit_pipeline_result` rule: `_`-prefixed keys are metadata, `reasoning`
    is the per-field trace artifact — never verification input)."""
    return {
        k: v
        for k, v in (extracted or {}).items()
        if not str(k).startswith("_") and k != "reasoning"
    }


def judge_verify_node(state: DocumentState) -> dict[str, Any]:
    """KANBAN-063 (Lane B): in-pipeline completeness judge.

    Reuses the offline-battle-tested ``CompletenessJudge`` rubric against the
    specialist's extraction. Verdict lands on state (`complete` proceeds,
    `partial`/`incomplete` go to the arbiter, hard failure escalates
    fail-safe). Gated by `_judge_gate` — most documents never see this call.
    """
    from agents.judge import CompletenessJudge
    from llm.retry import is_transient_error

    if not judge_gate(state):
        logger.info(
            "judge_skipped",
            doc_id=state.get("doc_id"),
            extraction_confidence=state.get("extraction_confidence"),
        )
        return {"judge_verdict": "skipped", "transient_error": False}

    doc_type = state.get("doc_type") or ""
    extracted = _clean_fields_for_judge(state.get("extracted_data"))
    doc_text = state.get("doc_text", "")
    try:
        judge = CompletenessJudge()
        result = judge.judge_completeness(
            doc_type=doc_type, extracted=extracted, doc_text=doc_text
        )
    except Exception as exc:
        if is_transient_error(exc):
            transient = state.get("transient_retries_judge_verify", 0) + 1
            logger.warning(
                "judge_transient_error",
                doc_id=state.get("doc_id"),
                error=str(exc)[:300],
                transient_retries=transient,
            )
            return {
                "transient_error": True,
                "transient_retries_judge_verify": transient,
                "error_message": f"transient provider error: {str(exc)[:200]}",
            }
        # Judge hard-failed AFTER the gate flagged this document as needing
        # scrutiny — escalate rather than archive unverified.
        logger.exception("judge_failed", doc_id=state.get("doc_id"))
        return {
            "judge_verdict": "judge_error",
            "escalation_reason": f"judge verification failed ({type(exc).__name__}) — routing to human review",
            "transient_error": False,
        }

    label = result.get("completeness_label", "incomplete")
    score = float(result.get("completeness", 0.0))
    findings = [str(result.get("reasoning", ""))] if result.get("reasoning") else []
    logger.info(
        "judge_verified",
        doc_id=state.get("doc_id"),
        label=label,
        score=score,
    )
    try:
        from observability.scores import emit_in_pipeline_judge_scores

        emit_in_pipeline_judge_scores(
            {
                "judge_verdict": label,
                "judge_score": score,
                "judge_findings": findings,
                "trace_id": state.get("trace_id"),
            }
        )
    except Exception:
        logger.exception("in_pipeline_judge_score_failed", doc_id=state.get("doc_id"))
    return {
        "judge_verdict": label,
        "judge_score": score,
        "judge_findings": findings,
        "judge_pass_count": int(state.get("judge_pass_count") or 0) + 1,
        "transient_error": False,
    }


def arbiter_node(state: DocumentState) -> dict[str, Any]:
    """KANBAN-063 (Lane B): arbitration on a failed judge verdict.

    Bounded decisions only (accept_with_caveats / retry with fix-list up to
    ``arbiter_retry_max`` / human_review). Any arbiter failure escalates
    fail-safe. Decisions are audited immediately and later copied onto the
    terminal manifest (archive or review/failed).
    """
    from agents.arbiter import ArbiterAgent
    from llm.retry import is_transient_error

    try:
        arbiter = ArbiterAgent()
        result = arbiter.arbitrate(
            doc_type=state.get("doc_type") or "",
            extracted=_clean_fields_for_judge(state.get("extracted_data")),
            judge_verdict=str(state.get("judge_verdict") or ""),
            judge_findings=state.get("judge_findings") or [],
            judge_score=state.get("judge_score"),
        )
    except Exception as exc:
        if is_transient_error(exc):
            transient = state.get("transient_retries_arbiter", 0) + 1
            logger.warning(
                "arbiter_transient_error",
                doc_id=state.get("doc_id"),
                error=str(exc)[:300],
                transient_retries=transient,
            )
            return {
                "transient_error": True,
                "transient_retries_arbiter": transient,
                "error_message": f"transient provider error: {str(exc)[:200]}",
                "escalation_reason": "transient provider error during arbitration",
            }
        logger.exception("arbiter_failed", doc_id=state.get("doc_id"))
        return {
            "arbiter_decision": "human_review",
            "arbiter_reasoning": f"arbiter failed ({type(exc).__name__})",
            "escalation_reason": f"arbitration failed ({type(exc).__name__}) — routing to human review",
            "transient_error": False,
        }

    decision = result.get("decision")
    updates: dict[str, Any] = {
        "arbiter_decision": decision,
        "arbiter_reasoning": str(result.get("reasoning", "")),
        "arbiter_handoff": str(result.get("handoff_summary", "")),
        "transient_error": False,
    }
    if decision == "retry_extraction":
        updates["arbiter_retry_count"] = state.get("arbiter_retry_count", 0) + 1
        updates["arbiter_fields_to_fix"] = list(result.get("fields_to_fix") or [])
        updates["escalation_reason"] = (
            f"arbiter ordered re-extraction: {result.get('handoff_summary', '')[:400]}"
        )
    elif decision == "accept_with_caveats":
        updates["escalation_reason"] = (
            f"arbiter accepted with caveats: {result.get('reasoning', '')[:400]}"
        )
    else:  # human_review
        updates["escalation_reason"] = (
            f"arbiter escalated to human review: {result.get('handoff_summary', '')[:400]}"
        )
    logger.info(
        "arbiter_decided",
        doc_id=state.get("doc_id"),
        decision=decision,
        retry_count=updates.get("arbiter_retry_count", state.get("arbiter_retry_count", 0)),
    )
    _emit_stage_audit(
        {**state, **updates},
        "arbiter_decided",
        actor="arbiter",
        detail={
            "decision": decision,
            "reasoning": updates.get("arbiter_reasoning"),
            "handoff": updates.get("arbiter_handoff"),
            "fields_to_fix": updates.get("arbiter_fields_to_fix"),
            "arbiter_retry_count": updates.get(
                "arbiter_retry_count", state.get("arbiter_retry_count", 0)
            ),
        },
    )
    return updates


def _normalize_review_decision(value) -> str | None:
    """Map a Command(resume=...) payload to approved/rejected, or None if invalid."""
    if value in ("approved", "rejected"):
        return value
    if value is True:
        return "approved"
    if value is False:
        return "rejected"
    if isinstance(value, str):
        token = value.strip().lower()
        if token in ("approved", "approve", "yes"):
            return "approved"
        if token in ("rejected", "reject", "no"):
            return "rejected"
        return None
    if isinstance(value, dict):
        inner = value.get("decision", value.get("approved"))
        if inner is value:
            return None
        return _normalize_review_decision(inner)
    return None


def _thread_is_interrupted(graph, thread_id: str) -> bool:
    """True when ``thread_id`` has a paused ``interrupt()`` checkpoint."""
    if not thread_id or graph is None:
        return False
    try:
        snap = graph.get_state({"configurable": {"thread_id": thread_id}})
    except Exception:
        return False
    if snap is None:
        return False
    if getattr(snap, "interrupts", None):
        return True
    for task in getattr(snap, "tasks", None) or ():
        if getattr(task, "interrupts", None):
            return True
    nxt = getattr(snap, "next", None) or ()
    return "human_review" in nxt


def _result_is_interrupted(result) -> bool:
    if not isinstance(result, dict):
        return False
    if result.get("__interrupt__"):
        return True
    return False


def _paused_review_result(result, initial_state, thread_id: str, trace_id: str) -> dict[str, Any]:
    """Normalize an ``interrupt()`` pause into a review-bin pipeline result.

    The node has already parked the file + catalog; its return value is not
    applied until resume, so stage/review_decision must be synthesized here.
    """
    out = (
        {k: v for k, v in result.items() if k != "__interrupt__"}
        if isinstance(result, dict)
        else dict(initial_state)
    )
    out["stage"] = PipelineStage.REVIEW.value
    out["review_decision"] = out.get("review_decision") or "pending_review"
    out["checkpoint_thread_id"] = thread_id
    if trace_id and not out.get("trace_id"):
        out["trace_id"] = trace_id
    doc_id = out.get("doc_id") or initial_state.get("doc_id")
    if doc_id:
        try:
            from pipeline.bins import load_manifest, save_manifest

            manifest = load_manifest(doc_id)
            if manifest is not None:
                manifest.checkpoint_thread_id = thread_id
                manifest.stage = PipelineStage.REVIEW
                if not manifest.review_decision:
                    manifest.review_decision = "pending_review"
                manifest.touch()
                save_manifest(manifest)
        except Exception:
            logger.exception("checkpoint_thread_persist_failed", doc_id=doc_id)
    logger.info("pipeline_paused_for_review", doc_id=doc_id, thread_id=thread_id)
    return out


def human_review_node(state: DocumentState) -> dict[str, Any]:
    """Park the document for a human, then ``interrupt()`` until they decide.

    Side effects before ``interrupt()`` are idempotent (park upsert + catalog
    upsert + manifest overwrite) because LangGraph restarts this node from
    the beginning on ``Command(resume=...)``.
    """
    from pipeline.bins import park_for_review
    from pipeline.reconsideration import collect_review_causes, format_causes

    file_path_str = state.get("file_path", "")
    esc_reason = state.get("escalation_reason", "Unknown reason")
    doc_id = state.get("doc_id", "")
    thread_id = state.get("checkpoint_thread_id") or ""

    logger.info("human_review_required", doc_id=doc_id, reason=esc_reason)

    causes = collect_review_causes(state)
    cause_line = format_causes(causes)
    if cause_line:
        if esc_reason and cause_line not in str(esc_reason):
            esc_reason = f"{esc_reason}; {cause_line}"
        else:
            esc_reason = cause_line

    dest: Path | None = None
    newly_parked = False
    if file_path_str:
        manifest = DocumentManifest(
            doc_id=doc_id,
            matter_id=state.get("matter_id", "DEFAULT"),
            original_filename=state.get("original_filename", ""),
            stage=PipelineStage.REVIEW,
            doc_type=state.get("doc_type"),
            contract_subtype=state.get("contract_subtype"),
            doc_subclass=state.get("doc_subclass"),
            classification_confidence=state.get("classification_confidence"),
            extracted_data=state.get("extracted_data"),
            extraction_confidence=state.get("extraction_confidence"),
            escalation_reason=esc_reason,
            trace_id=state.get("trace_id"),
            classification_attempts=state.get("classification_attempts", 0),
            extraction_attempts=state.get("extraction_attempts", 0),
            review_decision="pending_review",
            checkpoint_thread_id=thread_id or None,
            intake=state.get("intake_meta") or None,
            **_lane_b_manifest_fields(state),
        )
        dest, newly_parked = park_for_review(Path(file_path_str), manifest)
        _catalog_upsert(
            {
                "doc_id": doc_id,
                "matter_id": state.get("matter_id", "DEFAULT"),
                "original_filename": state.get("original_filename", ""),
                "doc_type": state.get("doc_type"),
                "contract_subtype": state.get("contract_subtype"),
                "doc_subclass": state.get("doc_subclass"),
                "classification_confidence": state.get("classification_confidence"),
                "extraction_confidence": state.get("extraction_confidence"),
                "extracted_data": {
                    **(state.get("extracted_data") or {} if isinstance(state.get("extracted_data"), dict) else {}),
                    "_lane_b": _lane_b_manifest_fields(state),
                },
                "escalation_reason": esc_reason,
                "trace_id": state.get("trace_id"),
                "stage": PipelineStage.REVIEW.value,
            },
            stage=PipelineStage.REVIEW.value,
        )

    parked = {
        "stage": PipelineStage.REVIEW.value,
        "escalation_reason": esc_reason,
        "review_causes": causes,
        "review_decision": "pending_review",
        "checkpoint_thread_id": thread_id or None,
        "file_path": str(dest) if dest is not None else file_path_str,
    }
    if newly_parked:
        _emit_stage_audit(
            {**state, **parked},
            "routed_to_review",
            actor="pipeline",
            detail={
                "reason": esc_reason,
                **{k: v for k, v in _lane_b_manifest_fields(state).items() if v not in (None, [], 0)},
            },
        )

    # Completion echo (HUB-037): a Gmail-intake document parked for review
    # still reports its outcome on the source email thread — dispatched AFTER
    # the routed_to_review audit entry so the echo's audit trail is complete.
    from pipeline.gmail_intake import dispatch_intake_echo

    dispatch_intake_echo(manifest.model_dump(mode="json"))

    # Relations clerk (HUB-040): the post-archive association pass — off the
    # document path (daemon thread), advisory, fail-soft.
    from pipeline.relations import dispatch_relations_scan

    dispatch_relations_scan(manifest.model_dump(mode="json"))

    payload = {
        "action": "human_review",
        "doc_id": doc_id or "",
        "matter_id": state.get("matter_id") or "DEFAULT",
        "filename": state.get("original_filename") or "",
        "doc_type": state.get("doc_type"),
        "escalation_reason": esc_reason or "",
        "review_causes": list(causes or []),
        "classification_confidence": state.get("classification_confidence"),
        "extraction_confidence": state.get("extraction_confidence"),
    }
    decision = None
    while decision is None:
        raw = interrupt(payload)
        decision = _normalize_review_decision(raw)
        if decision is None:
            payload = {
                **payload,
                "action": "human_review_invalid",
                "detail": "decision must be approved or rejected",
            }

    notes = ""
    if isinstance(raw, dict):
        notes = str(raw.get("notes") or "")

    if decision == "approved":
        logger.info("human_review_approved", doc_id=doc_id)
        return {
            "stage": PipelineStage.CLASSIFIED.value,
            "escalation_reason": None,
            "review_causes": causes,
            "review_decision": "approved",
            "extracted_data": None,
            "extraction_confidence": None,
            "extraction_attempts": 0,
            "classification_attempts": 0,
            # Fair Lane B budget on the corrected run; keep last arbiter_* /
            # judge_* fields on state for the eventual archive trail.
            "arbiter_retry_count": 0,
            "judge_pass_count": 0,
            "resume_extraction": True,
            "conflict_detected": False,
            "conflict_details": [],
            "file_path": str(dest) if dest is not None else file_path_str,
            "checkpoint_thread_id": thread_id or None,
            "transient_error": False,
        }

    logger.info("human_review_rejected", doc_id=doc_id, notes=notes[:200] or None)
    failed_path = str(dest) if dest is not None else file_path_str
    if dest is not None and dest.exists():
        failed_path = str(move_to_failed(dest))
    _catalog_upsert(
        {
            "doc_id": doc_id,
            "matter_id": state.get("matter_id", "DEFAULT"),
            "original_filename": state.get("original_filename", ""),
            "doc_type": state.get("doc_type"),
            "stage": PipelineStage.FAILED.value,
            "escalation_reason": esc_reason,
            "trace_id": state.get("trace_id"),
        },
        stage=PipelineStage.FAILED.value,
    )
    if doc_id:
        try:
            from pipeline.bins import load_manifest, save_manifest

            manifest = load_manifest(doc_id)
            if manifest is not None:
                manifest.stage = PipelineStage.FAILED
                manifest.review_decision = "rejected"
                manifest.touch()
                save_manifest(manifest)
        except Exception:
            logger.exception("review_reject_manifest_failed", doc_id=doc_id)
    _maybe_export_warehouse(doc_id)
    return {
        "stage": PipelineStage.FAILED.value,
        "escalation_reason": esc_reason,
        "review_causes": causes,
        "review_decision": "rejected",
        "file_path": failed_path,
        "checkpoint_thread_id": thread_id or None,
    }


def boss_escalation_node(state: DocumentState) -> dict[str, Any]:
    from agents.boss import BossAgent

    boss = BossAgent()
    manifest_data = {
        "doc_id": state.get("doc_id"),
        "doc_type": state.get("doc_type"),
        "classification_confidence": state.get("classification_confidence"),
        "extraction_confidence": state.get("extraction_confidence"),
        "extracted_data": state.get("extracted_data"),
        "escalation_reason": state.get("escalation_reason"),
    }

    # Give the Boss the archived matter context it adjudicates against, so its
    # decision is grounded in the actual conflicting records (it receives only
    # the doc manifest otherwise). Best-effort: empty context when unavailable.
    matter_context = _fetch_matter_context(state)

    # L-10: the Boss node was unguarded — a provider outage raised out of
    # adjudicate and the run's catch-all sent a well-classified/extracted
    # document to the FAILED bin (opposite of the Boss contract, which is to
    # route to human review when in doubt). Mirror the other nodes: transient
    # errors surface as a transient_error state (routing retries), and any
    # other exception defaults review_decision="review" so the document goes
    # to human review instead of being failed.
    try:
        result = boss.adjudicate(manifest_data, matter_context=matter_context)
        decision = result.get("decision", "review")
        reasoning = result.get("reasoning", "")
    except Exception as exc:
        from llm.retry import is_transient_error

        if is_transient_error(exc):
            transient = state.get("transient_retries_boss_escalation", 0) + 1
            logger.warning(
                "boss_transient_error",
                doc_id=state.get("doc_id"),
                error=type(exc).__name__,
                transient_retries=transient,
            )
            return {
                "transient_error": True,
                "transient_retries_boss_escalation": transient,
                "error_message": f"transient provider error: {str(exc)[:200]}",
                "escalation_reason": "transient provider error during boss adjudication",
            }
        logger.warning("boss_failed_defaulting_to_review", doc_id=state.get("doc_id"),
                       error=type(exc).__name__, exc_info=True)
        decision = "review"
        reasoning = f"Boss unavailable ({type(exc).__name__}) — defaulting to human review"

    logger.info(
        "boss_decision",
        doc_id=state.get("doc_id"),
        decision=decision,
        context_records=len(matter_context),
    )
    result = {
        "review_decision": decision,
        "escalation_reason": f"Boss: {reasoning}",
        "transient_error": False,
    }
    # A-1: Boss adjudication is part of the compliance record ("who decided
    # what, based on what").
    _emit_stage_audit({**state, **result}, "boss_adjudicated", actor="boss",
                      detail={"decision": decision, "reasoning": reasoning,
                              "context_records": len(matter_context)})
    return result


def compile_report_node(state: DocumentState) -> dict[str, Any]:
    """Procedural matter-record assembly (reporter LLM retired)."""
    from agents.reporter import compile_matter_record

    extracted = state.get("extracted_data") or {}
    if not isinstance(extracted, dict):
        extracted = {}
    manifest_data = {
        "doc_id": state.get("doc_id"),
        "matter_id": state.get("matter_id"),
        "doc_type": state.get("doc_type"),
        "contract_subtype": state.get("contract_subtype"),
        "doc_subclass": state.get("doc_subclass"),
        "classification_confidence": state.get("classification_confidence"),
        "extraction_confidence": state.get("extraction_confidence"),
        "extracted_data": extracted,
        "arbiter_decision": state.get("arbiter_decision"),
        "arbiter_reasoning": state.get("arbiter_reasoning"),
        "arbiter_handoff": state.get("arbiter_handoff"),
    }

    try:
        report = compile_matter_record(manifest_data)
    except Exception as exc:
        logger.exception("report_compile_failed", doc_id=state.get("doc_id"))
        report = {
            "summary": (
                f"(report unavailable: {type(exc).__name__} — extracted fields preserved)"
            ),
            "doc_type": state.get("doc_type"),
            "contract_subtype": state.get("contract_subtype"),
            "doc_subclass": state.get("doc_subclass"),
            "extracted_data": {
                k: v for k, v in extracted.items() if k not in ("confidence", "reasoning")
            },
            "classification_confidence": state.get("classification_confidence"),
            "extraction_confidence": state.get("extraction_confidence"),
            "error": True,
        }

    logger.info("report_compiled", doc_id=state.get("doc_id"), procedural=True)
    return {
        "extracted_data": {
            **extracted,
            "_report": report,
        },
        "report_error": bool(report.get("error")),
    }


def _catalog_upsert(state: dict, *, stage: str | None = None, update_only: bool = False) -> None:
    """Best-effort write/update the catalog record for a document state.

    The catalog is the durable conveyor position: every stage transition that
    moves a document (processing → review/failed/archived) upserts the row so
    `/ops/status`, stuck-doc detection, matter listings, and error rates reflect
    reality even when the run ends in review or fails. Never raises.
    """
    try:
        from schemas.matter import Matter
        from storage.catalog import write_document_record as _write_doc, write_matter_record as _write_matter

        doc_id = state.get("doc_id") or ""
        matter_id = state.get("matter_id") or "DEFAULT"
        if not doc_id:
            return

        doc_record = {
            "doc_id": doc_id,
            "matter_id": matter_id,
            "original_filename": state.get("original_filename", ""),
            "doc_type": state.get("doc_type", "unknown"),
            "contract_subtype": state.get("contract_subtype"),
            "doc_subclass": state.get("doc_subclass"),
            "stage": stage or state.get("stage", "cataloged"),
            "classification_confidence": state.get("classification_confidence"),
            "extraction_confidence": state.get("extraction_confidence"),
            "extracted_data": state.get("extracted_data"),
            "escalation_reason": state.get("escalation_reason") or state.get("error_message"),
            "trace_id": state.get("trace_id"),
        }

        def _sync_write():
            async def _runner():
                if not update_only:
                    matter = Matter(
                        matter_id=matter_id,
                        name=matter_id,
                        client_name="auto-created",
                        practice_area="transactional",
                    )
                    await _write_matter(matter)
                await _write_doc(doc_record)

            _run_coro(_runner)

        _sync_write()
        logger.debug("catalog_upserted", doc_id=doc_id, stage=doc_record["stage"])
    except Exception:
        logger.exception("catalog_upsert_error")


def catalog_write_node(state: DocumentState) -> dict[str, Any]:
    doc_id = state.get("doc_id", "")
    logger.info("catalog_write", doc_id=doc_id)
    # The catalog_write node runs BEFORE archive: the doc_type/confidence/
    # extraction fields are final here, but the terminal stage is not yet known
    # (archive_node completes the move). Write with the current stage; the
    # archive node upserts stage=archived afterwards so the catalog never
    # permanently shows a document as merely "classified".
    _catalog_upsert(state, stage=state.get("stage", "classified"))
    return {}


def _lane_b_manifest_fields(state: Mapping[str, Any] | dict) -> dict[str, Any]:
    """Copy judge/arbiter state onto a DocumentManifest / catalog payload."""
    fields_to_fix = state.get("arbiter_fields_to_fix")
    findings = state.get("judge_findings")
    return {
        "arbiter_decision": state.get("arbiter_decision"),
        "arbiter_reasoning": state.get("arbiter_reasoning"),
        "arbiter_handoff": state.get("arbiter_handoff"),
        "arbiter_fields_to_fix": list(fields_to_fix) if fields_to_fix else None,
        "arbiter_retry_count": int(state.get("arbiter_retry_count") or 0),
        "judge_verdict": state.get("judge_verdict"),
        "judge_score": state.get("judge_score"),
        "judge_findings": list(findings) if findings else None,
    }


def archive_node(state: DocumentState) -> dict[str, Any]:
    manifest = DocumentManifest(
        doc_id=state.get("doc_id", ""),
        matter_id=state.get("matter_id", "DEFAULT"),
        original_filename=state.get("original_filename", ""),
        stage=PipelineStage.ARCHIVED,
        doc_type=state.get("doc_type", "unknown"),
        contract_subtype=state.get("contract_subtype"),
        doc_subclass=state.get("doc_subclass"),
        classification_confidence=state.get("classification_confidence"),
        extracted_data=state.get("extracted_data"),
        extraction_confidence=state.get("extraction_confidence"),
        trace_id=state.get("trace_id"),
        escalation_reason=state.get("escalation_reason"),
        review_decision=state.get("review_decision"),
        classification_attempts=state.get("classification_attempts", 0),
        extraction_attempts=state.get("extraction_attempts", 0),
        intake=state.get("intake_meta") or None,
        **_lane_b_manifest_fields(state),
    )

    file_path_str = state.get("file_path", "")
    if not file_path_str:
        logger.error("archive_no_file_path", doc_id=manifest.doc_id)
        return _finalize_aborted(dict(state), "No file path in state")

    file_path = Path(file_path_str)
    if not file_path.exists():
        processing_root = processing_dir()
        candidates = list(processing_root.rglob(state.get("original_filename", "*.txt")))
        if not candidates:
            review_candidate = review_dir() / state.get("original_filename", "")
            if review_candidate.is_file():
                candidates = [review_candidate]
        if candidates:
            file_path = candidates[0]
        else:
            logger.error("archive_file_not_found", doc_id=manifest.doc_id, path=file_path_str)
            return _finalize_aborted(
                dict(state), f"File not found: {file_path_str}"
            )

    from agents.archivist import archive_document
    # Hash-chained audit trail: the new entry must link to the previous entry
    # for THIS doc_id (review decisions, re-runs, resumes all append to the
    # same chain). A stale `""` would break `verify_chain` for any second
    # event on the same document.
    prev_audit_hash = _latest_audit_hash(manifest.doc_id)
    archive_path, audit_entry = archive_document(manifest, file_path, prev_audit_hash=prev_audit_hash)

    _write_audit_log(audit_entry)

    # Final conveyor position: the catalog record (created at intake/catalog
    # write) must show archived, not classified — archive is the terminal stage.
    _catalog_upsert(
        {
            "doc_id": manifest.doc_id,
            "matter_id": manifest.matter_id,
            "original_filename": manifest.original_filename,
            "doc_type": manifest.doc_type,
            "contract_subtype": manifest.contract_subtype,
            "doc_subclass": getattr(manifest, "doc_subclass", None),
            "classification_confidence": manifest.classification_confidence,
            "extraction_confidence": manifest.extraction_confidence,
            "extracted_data": manifest.extracted_data,
            "escalation_reason": manifest.escalation_reason,
            "trace_id": manifest.trace_id,
            "stage": PipelineStage.ARCHIVED.value,
        },
        stage=PipelineStage.ARCHIVED.value,
        update_only=True,
    )

    logger.info("pipeline_complete", doc_id=manifest.doc_id, archive=str(archive_path))
    # Completion echo (HUB-037): Gmail-intake documents get the outcome
    # (archive entry + audit chain) replied onto their source email thread.
    from pipeline.gmail_intake import dispatch_intake_echo

    dispatch_intake_echo(manifest.model_dump(mode="json"))

    # Relations clerk (HUB-040): the post-archive association pass — off the
    # document path (daemon thread), advisory, fail-soft.
    from pipeline.relations import dispatch_relations_scan

    dispatch_relations_scan(manifest.model_dump(mode="json"))
    _maybe_export_warehouse(manifest.doc_id)
    return {"stage": PipelineStage.ARCHIVED.value}


def _maybe_export_warehouse(doc_id: str) -> None:
    """Best-effort Parquet cold-store export after a terminal stage (A-warehouse)."""
    if not doc_id:
        return
    try:
        from storage.warehouse import export_document_to_warehouse

        export_document_to_warehouse(doc_id)
    except Exception:
        logger.exception("warehouse_export_hook_failed", doc_id=doc_id)


def _run_coro(coro):
    """Run a coroutine from a sync context: schedule it on the running loop
    when one exists (thread-safe), otherwise run a fresh loop.

    `asyncio.get_event_loop()` is deprecated when no loop is running, and
    graph nodes execute both from the watcher's daemon threads (no loop) and
    from the API's `asyncio.to_thread` (a running loop in another thread), so
    both branches are needed.
    """
    import asyncio

    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.run(coro())
    import concurrent.futures

    future = asyncio.run_coroutine_threadsafe(coro(), loop)
    return future.result(timeout=10)


def _file_sha256(path) -> str:
    """SHA-256 of a file's bytes (audit A-7). Best-effort; "" on failure."""
    import hashlib

    try:
        h = hashlib.sha256()
        with open(path, "rb") as fh:
            for chunk in iter(lambda: fh.read(1 << 20), b""):
                h.update(chunk)
        return h.hexdigest()
    except Exception:
        logger.warning("file_sha256_failed", file=str(path))
        return ""


def _file_size(path) -> int:
    try:
        return path.stat().st_size
    except Exception:
        return 0


def _latest_audit_hash(doc_id: str) -> str:
    """Best-effort fetch of the last entry_hash for this doc_id (the previous
    link of the hash chain). Returns "" when no entries exist yet (a fresh
    chain) or the DB is unavailable — a broken link can never be caused by
    this: `build_audit_entry` uses whatever we return as prev_hash, and
    `verify_chain` recomputes from timestamps."""
    if not doc_id:
        return ""
    try:
        from storage.audit_log import get_latest_audit_hash

        return _run_coro(lambda: get_latest_audit_hash(doc_id))
    except Exception:
        logger.exception("latest_audit_hash_fetch_failed", doc_id=doc_id)
        return ""


def _write_audit_log(entry):
    try:
        import asyncio
        from storage.audit_log import write_audit_entry

        async def _write():
            await write_audit_entry(entry)

        _run_coro(_write)
    except Exception:
        logger.exception("audit_log_write_error")


def _emit_stage_audit(state: dict, event: str, actor: str = "pipeline", detail: dict | None = None) -> None:
    """Append a hash-chained audit entry for a stage transition (audit A-1).

    Before this, only the archivist and review decisions produced audit
    entries — classify/extract/retries/guardrails/Boss/review-routing/failures
    existed only in state + Langfuse. This gives every stage transition a
    chained, durable record so the compliance log answers "what decisions were
    made about this document, and when".

    Best-effort by design (an audit-write failure must never fail the run),
    but unlike the old swallow, the failure is logged with the event name so
    an operator can detect a gap (AUDIT_GAP).
    """
    doc_id = state.get("doc_id")
    if not doc_id:
        return
    try:
        from schemas.audit import build_audit_entry

        entry = build_audit_entry(
            doc_id=doc_id,
            matter_id=state.get("matter_id", "DEFAULT"),
            event=event,
            actor=actor,
            detail={
                **(detail or {}),
                "stage": state.get("stage"),
                "doc_type": state.get("doc_type"),
                "contract_subtype": state.get("contract_subtype"),
            "doc_subclass": state.get("doc_subclass"),
                "classification_confidence": state.get("classification_confidence"),
                "extraction_confidence": state.get("extraction_confidence"),
                "escalation_reason": state.get("escalation_reason"),
                "run_attempt": state.get("run_attempt"),
            },
            prev_hash=_latest_audit_hash(doc_id),
        )
        _write_audit_log(entry)
    except Exception:
        logger.warning("audit_gap", event=event, doc_id=doc_id, exc_info=True)


def _persist_scores(state: dict, scores: dict):
    if not scores or not state.get("doc_id"):
        return
    try:
        from storage.catalog import update_document_scores

        async def _write():
            await update_document_scores(state["doc_id"], scores)

        _run_coro(_write)
    except Exception:
        logger.exception("scores_persist_error")


def _persist_provenance(state: dict, run_id: str | None = None, metrics: dict | None = None) -> None:
    """A-10: persist doc → run → model → prompt → cost → latency provenance.

    The audit found `documents.trace_id` empty 60/60 in the live DB and no
    columns carrying run/model/prompt/cost — with tracing off, the demanded
    provenance chain was unreconstructable. Best-effort like other catalog
    writes (never fails the run)."""
    if not state.get("doc_id"):
        return
    try:
        from storage.catalog import update_document_provenance

        async def _write():
            await update_document_provenance(
                state["doc_id"],
                run_id=run_id or state.get("run_id"),
                model=_resolved_models(state),
                prompt_version=_prompt_versions(state),
                cost_usd=metrics.get("estimated_cost_usd") if metrics else None,
                latency_s=metrics.get("run_duration_seconds") if metrics else None,
                file_sha256=state.get("file_sha256"),
            )

        _run_coro(_write)
    except Exception:
        logger.exception("provenance_persist_error", doc_id=state.get("doc_id"))


def _resolved_models(state: dict) -> str:
    """Comma-joined model names actually used by this run (best-effort)."""
    try:
        from pipeline.config import get_agent_config, get_doc_class

        names: list[str] = []
        seen: set[str] = set()

        def _add(agent: str) -> None:
            if not agent or agent in seen:
                return
            seen.add(agent)
            cfg = get_agent_config(agent) or {}
            model = cfg.get("model")
            if model:
                names.append(f"{agent}={model}")

        _add("sorter")
        row = get_doc_class(state.get("doc_type") or "") or {}
        _add(str(row.get("specialist") or ""))
        if state.get("review_verdict"):
            _add("sorter_reviewer")
        if state.get("judge_verdict") and state.get("judge_verdict") != "skipped":
            _add("judge")
        if state.get("arbiter_decision"):
            _add("arbiter")
        if state.get("review_decision") in ("approved", "review") and (
            state.get("escalation_reason") or ""
        ).startswith("Boss:"):
            _add("boss")
        _add("reporter")
        return ", ".join(names)
    except Exception:
        return ""


def _prompt_versions(state: dict) -> str:
    """Prompt versions bound during the run (best-effort; Langfuse-managed
    prompts link generations; vendored agents use their versioned prompts)."""
    try:
        from llm.prompts import _bound_prompt_versions

        versions = _bound_prompt_versions()
        if versions:
            return ", ".join(f"{k}={v}" for k, v in versions.items())
    except Exception:
        pass
    try:
        from langchain_agents.prompts import PROMPT_VERSIONS

        return ", ".join(f"{k}={v}" for k, v in PROMPT_VERSIONS.items())
    except Exception:
        return ""


def _bounded(fn):
    """Node wrapper enforcing the per-run hard cutoff: wall-clock deadline and
    cumulative output-token budget. Raises RunDeadlineExceeded /
    RunBudgetExceeded, which `_execute_run` catches and finalizes as an
    aborted run (failed bin + scores).

    L-5: also emits a progress heartbeat — touches the catalog row's
    `updated_at` at every node boundary so long retry storms / big-PDF
    transcription keep the document "alive" to stuck detection instead of
    looking stale while actively burning LLM calls."""
    from pipeline.limits import check_run_deadline, check_token_budget

    @functools.wraps(fn)
    def wrapper(state):
        check_run_deadline(state.get("run_deadline"))
        check_token_budget()
        _touch_heartbeat(state)
        return fn(state)

    return wrapper


def _touch_heartbeat(state: dict) -> None:
    """L-5: best-effort `updated_at` bump for the document row (progress
    heartbeat). Stuck detection keys on updated_at; without this, a document
    in a long retry storm or PDF transcription looks stuck even though it is
    actively processing."""
    doc_id = state.get("doc_id")
    if not doc_id:
        return
    try:
        from storage.catalog import touch_document_heartbeat

        _run_coro(lambda: touch_document_heartbeat(doc_id))
    except Exception:
        pass  # heartbeat is best-effort; never fails the node


def build_graph(checkpointer=None):
    if checkpointer is None:
        checkpointer = _build_checkpointer()

    # Fail fast: every live taxonomy class must have a specialist extract arm.
    # A missing arm used to silently stub-extract and retry the same miss.
    _build_specialist_dispatch()

    workflow = StateGraph(DocumentState)

    # Node names stay stable (best practice); per-run values go in metadata.
    # Every node is bounded: the run deadline and token budget are enforced at
    # each boundary so a stuck run is cut off as soon as its budget is spent.
    workflow.add_node("intake", traced_node("intake-document")(_bounded(intake_node)))
    workflow.add_node("classify", traced_node("classify-document")(_bounded(classify_node)))
    workflow.add_node("retry_classify", traced_node("classify-document")(_bounded(retry_classify_node)))
    # KANBAN-062 (Lane A): agent second opinion on exhausted medium-band
    # classifications — the [REVIEW] exception path from the target diagram.
    workflow.add_node("review_classify", traced_node("classify-document")(_bounded(review_classify_node)))
    workflow.add_node("extract", traced_node("extract-fields")(_bounded(extract_node)))
    workflow.add_node("retry_extract", traced_node("extract-fields")(_bounded(retry_extract_node)))
    # KANBAN-063 (Lane B): gated completeness verification + arbitration —
    # the Judge → [ARBITER] exception path from the target diagram.
    workflow.add_node("judge_verify", traced_node("judge-verify")(_bounded(judge_verify_node)))
    workflow.add_node("arbiter", traced_node("arbitrate-verdict")(_bounded(arbiter_node)))
    workflow.add_node("human_review", traced_node("route-for-review")(_bounded(human_review_node)))
    workflow.add_node("boss_escalation", traced_node("adjudicate-conflict")(_bounded(boss_escalation_node)))
    workflow.add_node("compile_report", traced_node("compile-report")(_bounded(compile_report_node)))
    workflow.add_node("catalog_write", traced_node("write-catalog")(_bounded(catalog_write_node)))
    workflow.add_node("archive", traced_node("archive-document")(_bounded(archive_node)))

    workflow.add_conditional_edges(START, entry_route, {
        "intake": "intake",
        "extract": "extract",
    })
    workflow.add_edge("intake", "classify")

    workflow.add_conditional_edges("classify", after_classify, {
        "classify": "classify",  # transient-error self-loop (same node, LLM-level retry)
        "retry_classify": "retry_classify",
        "review_classify": "review_classify",  # GT class miss even at 0.99
        "extract": "extract",
        "human_review": "human_review",
    })

    workflow.add_conditional_edges("retry_classify", after_retry_classify, {
        "retry_classify": "retry_classify",  # transient self-loop (own per-node budget)
        "review_classify": "review_classify",
        "extract": "extract",
        "human_review": "human_review",
    })

    # KANBAN-062 (Lane A): agent second opinion. High-confidence reviewer →
    # extract (label applied by the node); anything else → human review.
    workflow.add_conditional_edges("review_classify", after_review_classify, {
        "review_classify": "review_classify",  # transient self-loop (own per-node budget)
        "extract": "extract",
        "human_review": "human_review",
    })

    workflow.add_conditional_edges("extract", after_extraction_gated, {
        "extract": "extract",  # transient-error self-loop (same node, LLM-level retry)
        "retry_extract": "retry_extract",
        "compile_report": "compile_report",
        "judge_verify": "judge_verify",
        "human_review": "human_review",
        "boss_escalation": "boss_escalation",
    })

    workflow.add_conditional_edges("retry_extract", after_retry_extraction_gated, {
        "retry_extract": "retry_extract",  # transient self-loop (own per-node budget)
        "compile_report": "compile_report",
        "judge_verify": "judge_verify",
        "human_review": "human_review",
        "boss_escalation": "boss_escalation",
    })

    # KANBAN-063 (Lane B): judge verdict routing — complete/skipped proceeds;
    # partial/incomplete goes to the arbiter; hard failure escalates fail-safe.
    workflow.add_conditional_edges("judge_verify", after_judge, {
        "judge_verify": "judge_verify",  # transient self-loop (own per-node budget)
        "compile_report": "compile_report",
        "arbiter": "arbiter",
        "human_review": "human_review",
    })

    workflow.add_conditional_edges("arbiter", after_arbiter, {
        "arbiter": "arbiter",  # transient self-loop (own per-node budget)
        "compile_report": "compile_report",
        "retry_extract": "retry_extract",
        "human_review": "human_review",
    })

    workflow.add_conditional_edges("boss_escalation", after_boss, {
        "boss_escalation": "boss_escalation",  # transient self-loop (own per-node budget)
        "compile_report": "compile_report",
        "human_review": "human_review",
    })

    workflow.add_conditional_edges("human_review", after_human_review, {
        "extract": "extract",
        "failed": END,
    })

    workflow.add_conditional_edges("compile_report", after_report, {
        "catalog_write": "catalog_write",
        "human_review": "human_review",
    })
    workflow.add_edge("catalog_write", "archive")
    workflow.add_edge("archive", END)

    return workflow.compile(checkpointer=checkpointer)


def _existing_processing_doc_id(original_filename: str) -> str | None:
    """Find the doc_id of an in-flight manifest for this filename.

    A run that crashed after intake already saved a processing-stage manifest
    (and a catalog row); the abort path must reuse that doc_id so the failed
    manifest/catalog record supersede the same document instead of orphaning
    the intake manifest and minting a second identity.
    """
    if not original_filename:
        return None
    try:
        import json as _json
        from pipeline.bins import manifests_dir

        mdir = manifests_dir()
        if not mdir.exists():
            return None
        for mf in mdir.glob("*.json"):
            try:
                data = _json.loads(mf.read_text())
            except Exception:
                continue
            if data.get("original_filename") != original_filename:
                continue
            if data.get("stage") == PipelineStage.PROCESSING.value:
                return data.get("doc_id")
    except Exception:
        logger.exception("abort_doc_id_lookup_failed")
    return None


def _finalize_aborted(initial_state: dict, reason: str, *, failure_class: str | None = None) -> dict:
    """Turn a run that hit a hard limit (or crashed) into a failed result.

    Moves the file to the failed bin with a manifest noting the abort, and
    returns a result dict that still carries doc/attempt fields so the run is
    scored (run_aborted=1) and visible in the catalog instead of stranding in
    processing/. ``failure_class`` distinguishes timeout vs auth vs I/O.
    """
    from pipeline.bins import move_to_failed, save_manifest
    from schemas.manifest import DocumentManifest, PipelineStage

    state = dict(initial_state)
    # Reuse the intake manifest's doc_id when the run crashed after intake, so
    # the aborted manifest supersedes the processing manifest (same identity).
    # Passed explicitly — DocumentManifest would otherwise mint a fresh UUID.
    aborted_doc_id = state.get("doc_id") or _existing_processing_doc_id(
        state.get("original_filename", "")
    )
    manifest_kwargs = dict(
        matter_id=state.get("matter_id", "DEFAULT"),
        original_filename=state.get("original_filename", ""),
        stage=PipelineStage.FAILED,
        doc_type=state.get("doc_type"),
        contract_subtype=state.get("contract_subtype"),
        doc_subclass=state.get("doc_subclass"),
        classification_confidence=state.get("classification_confidence"),
        classification_attempts=state.get("classification_attempts", 0),
        extracted_data=state.get("extracted_data"),
        extraction_confidence=state.get("extraction_confidence"),
        extraction_attempts=state.get("extraction_attempts", 0),
        escalation_reason=f"run aborted: {reason}",
        trace_id=state.get("trace_id"),
        intake=state.get("intake_meta") or None,
    )
    if aborted_doc_id:
        manifest_kwargs["doc_id"] = aborted_doc_id
    manifest = DocumentManifest(**manifest_kwargs)
    state["doc_id"] = manifest.doc_id
    state["stage"] = PipelineStage.FAILED.value
    state["run_aborted"] = True
    state["error_message"] = f"run aborted: {reason}"
    if failure_class:
        state["failure_class"] = failure_class
        tagged = f"run aborted [{failure_class}]: {reason}"
        manifest.escalation_reason = tagged
        state["error_message"] = tagged
        state["escalation_reason"] = tagged

    file_path_str = state.get("file_path") or ""
    if file_path_str:
        try:
            move_to_failed(Path(file_path_str))
        except Exception:
            logger.exception("abort_move_to_failed_error", file=file_path_str)
    try:
        save_manifest(manifest)
    except Exception:
        logger.exception("abort_manifest_save_error", doc_id=manifest.doc_id)
    try:
        _write_catalog_record(state)
    except Exception:
        logger.exception("abort_catalog_write_error", doc_id=manifest.doc_id)
    # A-1: a failed/aborted run is a compliance-record event (previously only
    # archived/reviewed runs left audit entries).
    try:
        _emit_stage_audit(
            state,
            "run_aborted",
            actor="pipeline",
            detail={"reason": reason, "failure_class": failure_class or "unexpected"},
        )
    except Exception:
        logger.exception("abort_audit_write_error", doc_id=manifest.doc_id)
    # Completion echo (HUB-037): failed Gmail-intake runs report on-thread
    # too — AFTER the run_aborted audit entry so the trail is complete.
    from pipeline.gmail_intake import dispatch_intake_echo

    dispatch_intake_echo(manifest.model_dump(mode="json"))

    # Relations clerk (HUB-040): the post-archive association pass — off the
    # document path (daemon thread), advisory, fail-soft.
    from pipeline.relations import dispatch_relations_scan

    dispatch_relations_scan(manifest.model_dump(mode="json"))
    _maybe_export_warehouse(manifest.doc_id)
    return state


def _write_catalog_record(state: dict):
    """Persist a minimal catalog record (used for aborted runs that never
    reach the catalog_write node, so they show up in compare_runs)."""
    import asyncio
    from storage.catalog import write_document_record

    doc_record = {
        "doc_id": state.get("doc_id", ""),
        "matter_id": state.get("matter_id", "DEFAULT"),
        "original_filename": state.get("original_filename", ""),
        "doc_type": state.get("doc_type", "unknown"),
        "stage": state.get("stage", "failed"),
        "classification_confidence": state.get("classification_confidence"),
        "extraction_confidence": state.get("extraction_confidence"),
        "extracted_data": state.get("extracted_data"),
        "escalation_reason": state.get("escalation_reason") or state.get("error_message"),
        "trace_id": state.get("trace_id"),
    }

    async def _write():
        await write_document_record(doc_record)

    _run_coro(_write)


# Cap for the judge-visible document text in the pipeline-result generation.
# Long contracts (e.g. MAUD merger agreements, ~800KB) would blow up the single
# cumulative judge call otherwise.
PIPELINE_RESULT_TEXT_LIMIT = 100_000


def _public_ground_truth(ground_truth: dict | None) -> dict:
    """Trace-safe ground truth (no expected_fields payloads / document text).

    The-Mailroom eval reads ``input.ground_truth.expected_hf_class`` and the
    flattened metadata keys. ``expected_fields`` stays off the trace.
    """
    if not ground_truth:
        return {}
    skip = {"expected_fields"}
    return {
        key: value
        for key, value in ground_truth.items()
        if key not in skip and value not in (None, "")
    }


def _emit_pipeline_result(root, result: dict, state: dict, judge_required: bool | None = None) -> None:
    """Emit the single `pipeline-result` generation observation per document.

    This is the ONLY observation the live evaluation rule matches
    (`scripts/sync_evaluators.py`, `mailroom-pipeline-rule`), so exactly one
    cumulative LLM-as-a-Judge call scores each document — classification
    correctness + extraction correctness + completeness in one pass. Output is
    the curated pipeline result. The input depends on the run mode:

    - Grounded runs (ground_truth carries `expected_fields`, i.e. pilot runs):
      the judge input is a labeled, pretty-printed expected-fields block. The
      output contains the labeled pipeline result and extracted fields. The
      full document text is NOT sent, which cuts the per-doc judge input from
      up to 100k chars (~25k tokens) to ~1-3k chars. The expected values ARE
      the ground truth, so the document text adds no verification power.
    - Live runs (no ground truth): the judge gets the (truncated) document
      text so it can verify grounding by rubric alone.

    Judge gating (issues #4/#5): for grounded runs the deterministic
    field-type-aware scorer (`observability/field_scoring.py`) runs first.
    When its verdict is unambiguous — every scored field clearly correct
    (above the ambiguous band) OR clearly wrong (below it) — `judge_required`
    is False and the generation is NOT emitted, so neither evaluator rule
    fires and two judge LLM calls are saved per document. When any field lands
    in the ambiguous band (0.5-0.85), or there is no ground truth at all
    (`judge_required is None`, live runs), the generation is emitted as
    before.

    No-ops when tracing is disabled (root None).

    A run that ends in `review` is NOT a final pipeline outcome: the document
    is awaiting (or has already received) a human decision, and the resumed
    run (if approved) re-archives under the same deterministic trace id. The
    generation is therefore suppressed for review-routed runs — the resumed
    run emits the single authoritative `pipeline-result`, so the evaluator
    fires exactly once per document trace instead of twice (once judged MISS
    for the review stage, once CORRECT after resume).
    """
    if root is None:
        return
    if result.get("stage") in ("review", "failed"):
        logger.info("pipeline_result_suppressed_non_terminal_stage", stage=result.get("stage"))
        return
    if judge_required is False:
        logger.info("pipeline_result_suppressed_deterministic_verdict")
        return
    import json

    ground_truth = state.get("ground_truth")
    grounded = bool(ground_truth and ground_truth.get("expected_fields"))
    extracted_data = result.get("extracted_data") or {}
    # `_report` is a derived catalog summary and may contain a full recursive
    # copy of the extraction. It is not part of any specialist schema and must
    # never be sent to the evaluator. `reasoning` is the v24+ per-field TRACE
    # artifact — it describes HOW values were found (never the values
    # themselves) and is likewise excluded from the evaluator's input.
    judge_extracted_data = {
        key: value for key, value in extracted_data.items()
        if not key.startswith("_") and key != "reasoning"
    }
    output = {
        "stage": result.get("stage"),
        "doc_type": result.get("doc_type"),
        "classification_confidence": result.get("classification_confidence"),
        "extraction_confidence": result.get("extraction_confidence"),
        "extracted_data": judge_extracted_data,
        "escalation_reason": result.get("escalation_reason"),
        "review_decision": result.get("review_decision"),
        "run_aborted": bool(result.get("run_aborted")),
        "error_message": result.get("error_message"),
    }
    if grounded:
        # Skip the document text entirely. The expected fields are the only
        # judge input; extracted fields are in the output, avoiding duplication.
        gen_input = (
            "GROUNDED EVALUATION INPUT\n"
            "The following fields are the literal expected values for this document.\n"
            "Compare them only with output.extracted_data for this same document.\n\n"
            "EXPECTED_FIELDS\n"
            "```json\n"
            f"{json.dumps(ground_truth['expected_fields'], ensure_ascii=False, indent=2)}\n"
            "```"
        )
        metadata = {"pipeline": "mailroom", "grounded": True, "input_format": "expected-fields-only"}
    else:
        doc_text = result.get("doc_text") or state.get("doc_text") or ""
        gen_input = doc_text[:PIPELINE_RESULT_TEXT_LIMIT]
        metadata = {
            "pipeline": "mailroom",
            "truncated": len(doc_text) > PIPELINE_RESULT_TEXT_LIMIT,
        }
    try:
        from observability.honest_gaps import honesty_trace_metadata

        honesty = honesty_trace_metadata(
            result.get("doc_type") or (ground_truth or {}).get("expected_doc_class"),
            extracted_data,
        )
        if honesty:
            metadata["suite_honesty"] = honesty
    except Exception:
        logger.debug("pipeline_result_honesty_attach_failed", exc_info=True)
    if ground_truth:
        # When the caller knows the expected outcome (pilot runs pass the
        # manifest ground truth), expose it here so the live evaluator can
        # decide a CORRECT/PARTIAL/MISS verdict against the ACTUAL truth
        # instead of judging by rubric alone. Expected fields already live in
        # the labeled input block, so do not duplicate them in the output.
        output["ground_truth"] = {
            key: value for key, value in ground_truth.items() if key != "expected_fields"
        }
    with observation(
        "pipeline-result",
        as_type="generation",
        input=gen_input,
        metadata=metadata,
    ) as gen:
        if gen is not None:
            gen.update(output=output)


def _execute_run(
    initial_state: DocumentState,
    seed: str,
    trace_input: dict,
    attempt: int = 0,
    source: str | None = None,
    ground_truth: dict | None = None,
    session_id: str | None = None,
    run_id: str | None = None,
    dataset: dict | None = None,
    invoke_input=None,
    thread_id: str | None = None,
) -> dict[str, Any]:
    """Shared execution scaffold: compile (or reuse) the graph, open the
    per-doc trace (one trace per document, deterministic id from `seed`),
    invoke, emit self-evident scores, persist them. Used by both
    `run_pipeline` and `resume_from_review`.

    Enforces the hard run cutoff: a wall-clock deadline and a cumulative
    output-token budget (see pipeline/limits.py). Aborted runs are finalized to
    the failed bin and still scored + persisted, so every run produces core
    metrics (duration, tokens, cost, call count) for cross-run evaluation.

    ``invoke_input`` + ``thread_id`` resume a paused ``interrupt()`` via
    ``Command(resume=...)``. An interrupt pause is a successful review park,
    not an aborted run.
    """
    import os
    from pipeline import limits
    from observability import tracing

    from observability import scores as pipeline_scores
    from langgraph.errors import GraphInterrupt

    # O-4: bind log correlation context so every log line in this run carries
    # doc/matter/run identifiers (merge_contextvars is wired in logging.py but
    # nothing ever bound anything). Unbound in finally.
    import structlog as _structlog

    _run_log = _structlog.contextvars
    _run_log.bind_contextvars(
        doc_id=initial_state.get("doc_id") or "",
        matter_id=initial_state.get("matter_id") or "",
        run_id=run_id or "",
        trace_id="",  # filled after the trace opens
    )

    graph = get_compiled_graph()

    # O-1: warm the score-config schema in a background thread (sticky-bounded
    # 10-min retry); never block the document path on Langfuse.
    pipeline_scores.warmup_score_configs(blocking=False)

    started_at = time.time()
    deadline = started_at + float(limits.get_deadline_seconds())
    limits.reset_run_usage()
    limits.set_run_deadline(deadline)
    thread_id = (
        thread_id
        or initial_state.get("checkpoint_thread_id")
        or f"{seed}-run{attempt}-{uuid.uuid4().hex[:8]}"
    )
    initial_state = {
        **initial_state,
        "run_deadline": deadline,
        "run_attempt": attempt,
        "checkpoint_thread_id": thread_id,
    }
    if ground_truth:
        # Expected outcome for this document (pilot runs pass the manifest
        # ground truth). Carried into the `pipeline-result` generation so the
        # live evaluator can render a CORRECT/PARTIAL/MISS verdict.
        initial_state["ground_truth"] = ground_truth
    payload = invoke_input if invoke_input is not None else initial_state

    # Environment resolution: per-context override (OBSERVABILITY_ENVIRONMENT,
    # set by entrypoints via pipeline.env.default_environment) wins; the
    # standard LANGFUSE_TRACING_ENVIRONMENT is the fallback; mock runs (no
    # observability) are labeled "mock"; everything else defaults to "live".
    environment = (
        os.environ.get("OBSERVABILITY_ENVIRONMENT")
        or os.environ.get("LANGFUSE_TRACING_ENVIRONMENT")
    )
    if not environment:
        if os.environ.get("OBSERVABILITY_PROVIDER", "auto") == "none":
            environment = "mock"
        else:
            environment = "live"

    # Mandatory tag taxonomy (see AGENTS.md "Mandatory: classify and tag every
    # logged run"): `mailroom` (always) + run-context tag matching the
    # environment (`pilot`/`live`/`misc`/`mock`) + attempt tag for re-runs +
    # source corpus tag for pilot/corpus runs.
    tags = ["mailroom", environment]
    if attempt:
        tags.append(f"run-{attempt}")
    if source:
        tags.append(f"source-{source}")
    try:
        from pipeline.docclass_mode import docclass_prompts_enabled

        if docclass_prompts_enabled():
            tags.append("docclass-prompts")
    except Exception:
        pass

    trace_metadata = {"pipeline": "mailroom", "run_deadline": deadline, "attempt": attempt}
    if source:
        trace_metadata["source"] = source
    if run_id:
        trace_metadata["run_id"] = run_id
    trace_metadata.update(dataset_trace_metadata(dataset))
    public_gt = _public_ground_truth(ground_truth)
    for key, value in public_gt.items():
        trace_metadata[key] = value
    # Honesty is a class property known once we have an expected/live type.
    # Tags stay immutable/upfront (mailroom/pilot/source-*); gap text goes in
    # metadata, never tags, and never a SCORE_CONFIG name that isn't registered.
    expected_class = public_gt.get("expected_doc_class") or public_gt.get("expected")
    if expected_class:
        try:
            from observability.honest_gaps import honesty_trace_metadata

            honesty = honesty_trace_metadata(str(expected_class))
            if honesty:
                trace_metadata["suite_honesty"] = honesty
        except Exception:
            logger.debug("suite_honesty_attach_failed", exc_info=True)

    # Native LangGraph RunnableConfig: thread-scoped attempt (a re-run of the
    # same document must not resume the previous run's checkpointed state —
    # pilot: correspondence_01's degraded second run inherited stale state)
    # plus run-level tags/metadata. LangGraph propagates these natively to any
    # callback/instrumentation attached to the run, so the graph-level run
    # carries the same classification dimensions as the Langfuse trace.
    config = {
        "configurable": {"thread_id": thread_id},
        "tags": tags,
        "metadata": trace_metadata,
    }

    with tracing.pipeline_trace(
        seed=seed,  # deterministic trace id -> correlates with our doc
        session_id=session_id or initial_state.get("matter_id") or "DEFAULT",  # groups documents of a matter/run
        name="document-pipeline",
        input=trace_input,
        metadata=trace_metadata,
        tags=tags,
        environment=environment,
        user_id=os.environ.get("MAILROOM_TRACE_USER_ID") or None,
        as_type="chain",
    ) as root:
        # Capture the trace id into the state so manifests, catalog records and
        # the returned result all carry it (the DB↔Langfuse correlation link).
        # get_trace_id() is only valid inside the trace block.
        state_trace_id = tracing.get_trace_id() or ""
        initial_state = {**initial_state, "trace_id": state_trace_id}
        if invoke_input is None:
            payload = initial_state
        _run_log.bind_contextvars(trace_id=state_trace_id)
        try:
            result = graph.invoke(payload, config)
        except GraphInterrupt:
            result = {**initial_state, "__interrupt__": True}
        except (limits.RunDeadlineExceeded, limits.RunBudgetExceeded) as exc:
            from pipeline.failures import classify_run_failure

            classified = classify_run_failure(exc)
            logger.warning(
                "run_aborted",
                doc_id=initial_state.get("doc_id"),
                reason=classified["reason"],
                failure_class=classified["failure_class"],
                detail=str(exc),
            )
            result = _finalize_aborted(
                initial_state,
                classified["reason"],
                failure_class=classified["failure_class"],
            )
        except Exception as exc:
            from pipeline.failures import classify_run_failure

            classified = classify_run_failure(exc)
            logger.exception(
                "run_crashed",
                doc_id=initial_state.get("doc_id"),
                failure_class=classified["failure_class"],
            )
            result = _finalize_aborted(
                initial_state,
                classified["reason"],
                failure_class=classified["failure_class"],
            )
        if _result_is_interrupted(result):
            result = _paused_review_result(result, initial_state, thread_id, state_trace_id)
        # Ensure the trace id survives into the final state (intake_node creates
        # the manifest with its own doc_id; the trace id must be attached even
        # when the graph never ran intake, e.g. aborted runs).
        if not result.get("trace_id"):
            result["trace_id"] = state_trace_id

        # Post-invoke scoring/emission is best-effort: the file may already be
        # archived/reviewed, so a failure here must never surface as a pipeline
        # failure to the watcher (the run itself succeeded).
        try:
            metrics = pipeline_scores.compute_run_metrics(result, started_at, time.time())
            score_values = pipeline_scores.emit_pipeline_scores(result, metrics)
            _persist_scores(result, score_values)
        except Exception:
            logger.exception("post_run_scoring_failed", doc_id=result.get("doc_id"))
        # A-10: persist end-to-end provenance (doc → run → trace → prompt →
        # model → cost) to the catalog row so the chain is reconstructable even
        # with tracing off (the live-DB audit found trace_id empty 60/60).
        try:
            _persist_provenance(result, run_id=run_id, metrics=metrics if "metrics" in dir() else None)
        except Exception:
            logger.exception("provenance_persist_failed", doc_id=result.get("doc_id"))
        # Deterministic field-type-aware scoring (issues #4/#5). Only grounded
        # runs have expected field values to compare against; the scorer gates
        # the LLM judge below: an unambiguous deterministic verdict skips the
        # `pipeline-result` generation entirely (2 evaluator calls saved).
        judge_required = None
        expected_fields = (initial_state.get("ground_truth") or {}).get("expected_fields")
        if expected_fields:
            from observability.field_scoring import get_field_types, score_extraction
            from observability.langfuse_field_scoring import score_and_log_extraction

            extracted = result.get("extracted_data") or {}
            doc_class = result.get("doc_type")
            expected_class = (initial_state.get("ground_truth") or {}).get("expected_doc_class")
            if doc_class and extracted:
                try:
                    from observability.extraction_gt import presence_expectations_from_ground_truth

                    presence = presence_expectations_from_ground_truth(
                        initial_state.get("ground_truth") or {},
                        doc_class,
                    )
                    field_result = score_and_log_extraction(
                        trace_id=tracing.get_trace_id(),
                        doc_class=doc_class,
                        field_types=get_field_types(doc_class),
                        predicted=extracted,
                        expected=expected_fields,
                        matter_id=initial_state.get("matter_id"),
                        doc_text=initial_state.get("doc_text"),
                        expected_class=expected_class,
                        presence_expectations=presence,
                    )
                    judge_required = field_result.needs_judge_review
                    # A wrong classification must ALWAYS reach the LLM judge:
                    # the deterministic field scorer compares the extraction
                    # against the EXPECTED class's fields, so a misfiled doc
                    # scores ~0 (below the ambiguous band) and would suppress
                    # the verdict for exactly the runs that need scrutiny.
                    if expected_class and doc_class != expected_class:
                        judge_required = True
                        logger.info(
                            "field_scoring_class_mismatch_forces_judge",
                            doc_id=result.get("doc_id"),
                            doc_class=doc_class,
                            expected_class=expected_class,
                        )
                    logger.info(
                        "field_scoring_computed",
                        doc_id=result.get("doc_id"),
                        overall=field_result.overall_score,
                        ambiguous_fields=field_result.ambiguous_fields,
                        judge_required=judge_required,
                    )
                except Exception:
                    logger.exception("field_scoring_failed")
        try:
            if root is not None:
                _emit_pipeline_result(root, result, initial_state, judge_required=judge_required)
                output = {
                    "stage": result.get("stage"),
                    "doc_type": result.get("doc_type"),
                    "classification_confidence": result.get("classification_confidence"),
                    "extraction_confidence": result.get("extraction_confidence"),
                    "run_aborted": bool(result.get("run_aborted")),
                    "error_message": result.get("error_message"),
                }
                if public_gt:
                    output["ground_truth"] = public_gt
                root.update(output=output)
        except Exception:
            logger.exception("pipeline_result_emission_failed", doc_id=result.get("doc_id"))
    # O-2: the flush must run in `finally` — hard failures raised out of the
    # graph invoke (before the old flush line) dropped every buffered trace
    # event. Bounded by the SDK's own send/retry; a blackholed backend is
    # surfaced via flush_health() counters instead of blocking forever.
    try:
        tracing.flush()
    except Exception:
        logger.warning("tracing_flush_error", doc_id=result.get("doc_id"), exc_info=True)
    _run_log.unbind_contextvars("doc_id", "matter_id", "run_id", "trace_id")  # O-4
    return result


def dataset_trace_metadata(dataset: dict | None) -> dict[str, str]:
    """§45 evaluation-trace identity (HUB-022 P0 residual): every evaluation
    trace records dataset_name, dataset_revision, taxonomy_version.
    matter_id rides as session_id and run_id is the simulation-run surrogate
    (already on the trace). Empty dict when unset — live (non-corpus) runs
    carry no dataset identity rather than a fabricated one.
    """
    if not isinstance(dataset, dict):
        return {}
    out: dict[str, str] = {}
    for meta_key, src in (
        ("dataset_name", "name"),
        ("dataset_revision", "revision"),
        ("taxonomy_version", "taxonomy_version"),
    ):
        value = dataset.get(src)
        if value not in (None, ""):
            out[meta_key] = str(value)
    return out


def run_pipeline(
    file_path: Path,
    matter_id: str = "DEFAULT",
    attempt: int = 0,
    source: str | None = None,
    ground_truth: dict | None = None,
    session_id: str | None = None,
    run_id: str | None = None,
    dataset: dict | None = None,
    intake_meta: dict | None = None,
) -> dict[str, Any]:
    _ensure_dirs()

    initial_state: DocumentState = {
        "doc_id": "",
        "matter_id": matter_id,
        "original_filename": file_path.name,
        "stage": "inbox",
        "doc_type": None,
        "classification_confidence": None,
        "classification_attempts": 0,
        "extracted_data": None,
        "extraction_confidence": None,
        "extraction_attempts": 0,
        "trace_id": None,
        "escalation_reason": None,
        "review_decision": None,
        "retry_count": 0,
        "conflict_detected": False,
        "file_path": str(file_path),
        "doc_text": "",
        "doc_pages": [],
        "error_message": None,
        "messages": [],
        "transient_error": False,
        "transient_retries_classify": 0,
        "transient_retries_extract": 0,
        "run_attempt": attempt,
    }
    if intake_meta:
        # Intake provenance (HUB-037): carried by state so every manifest
        # construction site (intake / review / archive / aborted) records it.
        initial_state["intake_meta"] = dict(intake_meta)

    # Attempt 0 keeps the bare filename stem as the deterministic trace seed
    # (backwards-compatible with ground-truth score ingestion in run_pilot.py);
    # subsequent attempts (e.g. scheduled re-processing) get a suffixed seed so
    # each run gets its own trace instead of merging into one misleading span.
    seed = file_path.stem if attempt <= 0 else f"{file_path.stem}-run{attempt}"

    trace_input = {"filename": file_path.name, "matter_id": matter_id, "attempt": attempt}
    public_gt = _public_ground_truth(ground_truth)
    if public_gt:
        # The-Mailroom / eval_pipeline.py prefer input.ground_truth (and the
        # matching metadata keys) over the pipeline-result generation.
        trace_input["ground_truth"] = public_gt

    return _execute_run(
        initial_state,
        seed=seed,
        attempt=attempt,
        source=source,
        ground_truth=ground_truth,
        session_id=session_id,
        run_id=run_id,
        dataset=dataset,
        trace_input=trace_input,
    )


def resume_from_review(manifest, review_file: Path, notes: str = "") -> dict[str, Any]:
    """Resume a human-approved review document with a FRESH extraction.

    Prefers LangGraph ``Command(resume=...)`` on the parked ``interrupt()``
    checkpoint (same process, MemorySaver still live). If the checkpoint is
    gone (process restart + MemorySaver), falls back to the stateless
    re-invoke: requeue from the review bin and start at extract under the
    original ``doc_id``.
    """
    from pipeline import limits
    from pipeline.bins import requeue_from_review, get_worker_id

    if not manifest.doc_type:
        raise ValueError(
            "Cannot resume: manifest has no classification; re-submit the document instead."
        )

    _ensure_dirs()
    thread_id = getattr(manifest, "checkpoint_thread_id", None) or ""
    graph = get_compiled_graph()
    if thread_id and _thread_is_interrupted(graph, thread_id):
        deadline = time.time() + float(limits.get_deadline_seconds())
        logger.info(
            "review_interrupt_resume",
            doc_id=manifest.doc_id,
            thread_id=thread_id,
        )
        return _execute_run(
            {
                "doc_id": manifest.doc_id,
                "matter_id": manifest.matter_id,
                "original_filename": manifest.original_filename,
                "doc_type": manifest.doc_type,
                "checkpoint_thread_id": thread_id,
                "review_decision": "approved",
            },
            seed=Path(manifest.original_filename).stem,
            attempt=0,
            thread_id=thread_id,
            invoke_input=Command(
                resume={"decision": "approved", "notes": notes or ""},
                update={"run_deadline": deadline},
            ),
            trace_input={
                "filename": manifest.original_filename,
                "matter_id": manifest.matter_id,
                "resumed": True,
                "interrupt_resume": True,
            },
        )

    worker_id = get_worker_id()
    queued = requeue_from_review(review_file, worker_id)
    doc_text, text_ok = _read_file_text(queued)
    from agents.intake import apply_intake

    doc_text, _intake_stats = apply_intake(doc_text, filename=queued.name)
    doc_pages = _render_doc_pages(queued)

    initial_state: DocumentState = {
        "doc_id": manifest.doc_id,
        "matter_id": manifest.matter_id,
        "original_filename": manifest.original_filename,
        "stage": PipelineStage.CLASSIFIED.value,
        "doc_type": manifest.doc_type,
        "contract_subtype": manifest.contract_subtype,
        "doc_subclass": getattr(manifest, "doc_subclass", None),
        "classification_confidence": manifest.classification_confidence,
        "classification_attempts": 0,
        "extracted_data": None,  # fresh extraction — never reuse the reviewed data
        "extraction_confidence": None,
        "extraction_attempts": 0,
        "trace_id": None,
        "escalation_reason": None,
        "review_decision": "approved",
        "retry_count": 0,
        "conflict_detected": False,
        "file_path": str(queued),
        "doc_text": doc_text,
        "doc_pages": doc_pages,
        "error_message": None if text_ok else f"Could not extract text from {queued.suffix} file",
        "messages": [],
        "resume_extraction": True,
        "transient_error": False,
        "transient_retries_classify": 0,
        "transient_retries_extract": 0,
        # Retain prior Lane B trail from the parked manifest for archive audit;
        # reset attempt counters so the corrected run gets a fair QA budget.
        "reviewer_doc_type": None,
        "reviewer_contract_subtype": None,
        "reviewer_doc_subclass": None,
        "reviewer_confidence": None,
        "review_verdict": None,
        "judge_verdict": getattr(manifest, "judge_verdict", None),
        "judge_score": getattr(manifest, "judge_score", None),
        "judge_findings": list(getattr(manifest, "judge_findings", None) or []),
        "judge_pass_count": 0,
        "arbiter_decision": getattr(manifest, "arbiter_decision", None),
        "arbiter_reasoning": getattr(manifest, "arbiter_reasoning", None),
        "arbiter_handoff": getattr(manifest, "arbiter_handoff", None),
        "arbiter_fields_to_fix": list(getattr(manifest, "arbiter_fields_to_fix", None) or []),
        "arbiter_retry_count": 0,
        "run_attempt": 0,
    }

    logger.info(
        "review_stateless_resume",
        doc_id=manifest.doc_id,
        reason="no interrupt checkpoint",
    )
    return _execute_run(
        initial_state,
        seed=f"{queued.stem}-resume-{manifest.doc_id[:8]}",
        attempt=0,
        trace_input={"filename": queued.name, "matter_id": manifest.matter_id, "resumed": True},
    )
