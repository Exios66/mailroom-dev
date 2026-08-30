"""P0.4 — .docx ingestion: text extraction (python-docx) instead of decoding
the zip binary as UTF-8, which previously produced garbage text for the
classifier. Unreadable .docx files return the standard unreadable marker with
ok=False so the pipeline routes them to review rather than misclassifying."""

from pathlib import Path

from graph.build_graph import _read_file_text


def _write_docx(path: Path, paragraphs: list[str]) -> None:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Party A"
    table.rows[0].cells[1].text = "Party B"
    doc.save(str(path))


class TestDocxIngestion:
    def test_docx_text_extracted_with_paragraphs_and_tables(self, temp_base_dir):
        path = temp_base_dir / "sample.docx"
        _write_docx(path, ["Master Services Agreement", "This agreement is governed by Delaware law."])

        text, ok = _read_file_text(path)

        assert ok is True
        assert "Master Services Agreement" in text
        assert "Delaware law" in text
        assert "Party A | Party B" in text

    def test_docx_not_decoded_as_binary_garbage(self, temp_base_dir):
        """Regression: before the fix, .docx decoded as UTF-8 produced zip
        gibberish (e.g. 'PK\\x03\\x04...')."""
        path = temp_base_dir / "sample.docx"
        _write_docx(path, ["Agreement"])

        text, ok = _read_file_text(path)

        assert ok is True
        assert "PK" not in text[:32]

    def test_invalid_docx_rejected_with_clear_error(self, temp_base_dir):
        path = temp_base_dir / "corrupt.docx"
        path.write_bytes(b"this is not a valid docx zip")

        text, ok = _read_file_text(path)

        assert ok is False
        assert "Unreadable file" in text

    def test_empty_docx_returns_empty(self, temp_base_dir):
        from docx import Document

        path = temp_base_dir / "empty.docx"
        Document().save(str(path))

        text, ok = _read_file_text(path)

        assert ok is False
        assert text == ""
